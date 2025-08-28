from django.shortcuts import render, redirect, get_object_or_404
from .models import User
from .forms import UserCreationForm, UserUpdateForm, ClassGroupForm
from django.contrib import messages
from .models import ClassGroup # Pastikan ClassGroup diimpor di bagian atas
from .forms import ClassGroupForm # Pastikan ClassGroupForm juga diimpor
import pandas as pd
from django.http import HttpResponse
from io import BytesIO
from .models import MataPelajaran # Tambahkan MataPelajaran ke impor
from .forms import MataPelajaranForm # Tambahkan MataPelajaranForm ke impor
from .models import Question, Exam
from .forms import QuestionForm
import json
from .models import Exam, ExamAttempt, Answer
from .forms import ExamForm, AddQuestionsToExamForm
from django.db import transaction
from .models import TeacherAssignment, TeacherEvaluation # Pastikan TeacherAssignment sudah diimpor
from .forms import EvaluationForm # Tambahkan EvaluationForm ke impor
from .models import Enrollment # Pastikan Enrollment diimpor
from io import BytesIO
from .forms import ProfileUpdateForm, CustomPasswordChangeForm
from django.contrib.auth import update_session_auth_hash
from django.db.models import Subquery, OuterRef
from .forms import EvaluationForm, TeacherProfileUpdateForm # Tambahkan form baru
from .forms import UserUpdateForm
from django.shortcuts import redirect
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.db.models import Subquery, OuterRef, Case, When, Value, CharField
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db.models import Subquery, OuterRef
from .forms import ClassForm, ClassGroupForm
from .models import GradingSettings
from .forms import GradingSettingsForm
from .models import MataPelajaran, Competency # Pastikan import sudah ada
from .forms import CompetencyForm # Pastikan import sudah ada
from .models import User, TeacherAssignment, ClassGroup, MataPelajaran, Competency, Score
from django.views.decorators.http import require_POST
from .models import MataPelajaran, Competency, AssessmentComponent
import docx, os, ast
from django.http import FileResponse
from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.db.models import Q, Subquery, OuterRef, Case, When, Value, CharField
from django.conf import settings
from django.core.paginator import Paginator 
from collections import defaultdict
from .forms import AddQuestionsToExamForm
from django.core.serializers.json import DjangoJSONEncoder
from django.contrib import messages
from .forms import ExamForm
from .models import Exam, Token, ExamAttempt
from django.contrib.auth.decorators import login_required
from django.forms import formset_factory
from django.db.models import Count, Q
from django.utils import timezone
from openpyxl.styles import Font, Alignment, Border, Side
import google.generativeai as genai
from django.http import JsonResponse
from django.forms import inlineformset_factory
from .models import User, ClassGroup, Ekstrakurikuler, Kehadiran, InfoSekolah, MataPelajaran, Rapor, Penilaian, KomponenNilai
from django.core.exceptions import ObjectDoesNotExist
from .forms import UserEditForm
from openpyxl import Workbook
from .models import MataPelajaran
from .forms import MataPelajaranForm
import openpyxl
from .forms import ThemeForm
from django.contrib.auth import get_user_model
import inspect
import re
import io
import uuid
from dashboard.models import ExamAttempt
from django.contrib.auth.models import User
from .models import ClassGroup, MataPelajaran
import logging
from django.contrib.auth.models import User
from django.views.decorators.csrf import csrf_exempt
from .forms import KomponenNilaiForm
import html
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
import secrets
from django.core.serializers.json import DjangoJSONEncoder
from .forms import StudentProfileForm

# Menyiapkan logger untuk mencatat error
logger = logging.getLogger(__name__)

def is_admin(user):
    """
    Mengecek apakah pengguna memiliki peran sebagai Guru, Walikelas, atau Admin
    berdasarkan field 'role' dari model User kustom Anda.
    """
    # Menggunakan konstanta dari model User Anda (user.Role) untuk perbandingan yang aman
    # Ini mengizinkan Guru, Walikelas, dan Admin untuk mengakses halaman yang dilindungi.
    return user.role in [User.Role.GURU, User.Role.WALIKELAS, User.Role.ADMIN]

User = get_user_model()
def halaman_utama(request):
    # Fungsi ini hanya akan merender (menampilkan) file home.html
    return render(request, 'dashboard/main/home.html')

@login_required
def user_management_view(request):
    """
    View yang telah diperbaiki untuk manajemen pengguna, menggabungkan semua fungsionalitas:
    - Menampilkan kelas untuk Siswa & Wali Kelas.
    - Menambah pengguna baru.
    - Hapus massal pengguna.
    - Filter berdasarkan peran.
    - Pencarian.
    - Paginasi.
    """
    
    # --- LOGIKA UNTUK PERMINTAAN POST (TAMBAH USER & HAPUS MASSAL) ---
    form = UserCreationForm()
    if request.method == 'POST':
        # Logika untuk Hapus Massal
        if 'delete_selected' in request.POST:
            user_ids_to_delete = request.POST.getlist('user_ids')
            if user_ids_to_delete:
                users_to_delete = User.objects.filter(pk__in=user_ids_to_delete).exclude(pk=request.user.pk)
                count = users_to_delete.count()
                users_to_delete.delete()
                messages.success(request, f'{count} pengguna berhasil dihapus.')
            return redirect('user_management')
        
        # Logika untuk Tambah Pengguna Baru
        else:
            form = UserCreationForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, 'Pengguna baru berhasil ditambahkan!')
                return redirect('user_management')

    # --- LOGIKA UNTUK PERMINTAAN GET (MENAMPILKAN DATA, FILTER, PENCARIAN) ---
    selected_role = request.GET.get('role', '')
    search_query = request.GET.get('search', '')

    # --- INI ADALAH KUNCI PERBAIKANNYA ---
    # Subquery untuk mendapatkan kelas siswa dari model Enrollment
    student_class = Enrollment.objects.filter(
        student=OuterRef('pk')
    ).values('class_group__name')[:1]

    # Subquery untuk mendapatkan kelas yang diampu wali kelas
    homeroom_class = ClassGroup.objects.filter(
        homeroom_teacher=OuterRef('pk')
    ).values('name')[:1]

    # Kueri dasar dengan anotasi yang benar untuk kolom 'class_name'
    users_query = User.objects.annotate(
        class_name=Case(
            When(role='WALIKELAS', then=Subquery(homeroom_class)),
            When(role='SISWA', then=Subquery(student_class)),
            default=Value(''),  # Gunakan string kosong, lebih mudah di template
            output_field=CharField()
        )
    ).order_by('first_name', 'last_name')

    # Terapkan filter berdasarkan peran
    if selected_role:
        users_query = users_query.filter(role=selected_role)
    
    # Terapkan filter pencarian
    if search_query:
        users_query = users_query.filter(
            Q(username__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query)
        )

    # Menambahkan Paginasi untuk performa lebih baik
    paginator = Paginator(users_query, 25) # Menampilkan 25 pengguna per halaman
    page_number = request.GET.get('page')
    users_page = paginator.get_page(page_number)

    context = {
        'form': form,
        'users': users_page, # Kirim objek halaman yang sudah dipaginasi
        'roles': User.Role.choices,
        'selected_role': selected_role,
        'search_query': search_query,
    }
    return render(request, 'dashboard/user_management.html', context)

def user_update_view(request, pk):
    """
    View untuk mengedit data pengguna.
    """
    user = get_object_or_404(User, pk=pk)
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, request.FILES, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, f'Data pengguna {user.username} berhasil diperbarui!')
            return redirect('user_management')
    else:
        form = UserUpdateForm(instance=user)

    context = {
        'form': form,
        'page_title': f'Edit Pengguna: {user.username}'
    }
    return render(request, 'dashboard/user_form.html', context)

def user_delete_view(request, pk):
    """
    View untuk menghapus pengguna.
    """
    user = get_object_or_404(User, pk=pk)
    if request.method == 'POST':
        username = user.username
        user.delete()
        messages.success(request, f'Pengguna {username} berhasil dihapus.')
        return redirect('user_management')

    context = {
        'user_to_delete': user
    }
    return render(request, 'dashboard/user_confirm_delete.html', context)

@login_required
def class_management_view(request):
    """
    View yang diperbaiki untuk menangani form dan filter jenjang dengan benar.
    """
    selected_level = request.GET.get('level', '')
    levels = ['SD', 'SMP', 'SMA']
    
    class_query = ClassGroup.objects.all().prefetch_related('students').order_by('name')
    if selected_level in levels:
        class_query = class_query.filter(level=selected_level)

    all_MataPelajarans = MataPelajaran.objects.all()

    if request.method == 'POST' and 'add_class' in request.POST:
        # Gunakan ClassForm yang sudah kita perbaiki
        form = ClassForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Kelas baru berhasil ditambahkan.')
            return redirect('class_management')
    else:
        form = ClassForm()
    
    total_students_count = User.objects.filter(role='SISWA').count()
    print(f"Jumlah siswa yang ditemukan: {total_students_count}") 

    context = {
        'form': form,
        'classes': class_query,
        'levels': levels,
        'selected_level': selected_level,
        'all_MataPelajarans': all_MataPelajarans,
        'total_students': total_students_count, # <-- TAMBAHKAN BARIS INI
    }
    return render(request, 'dashboard/class_management.html', context)

# =======================================================
# FUNGSI BARU UNTUK IMPOR/EKSPOR KELAS
# =======================================================
@login_required
def class_export_template(request):
    """
    Men-download file template Excel kosong untuk membuat kelas baru.
    Kolom yang dibutuhkan adalah 'name' dan 'level'.
    """
    data = {
        'name': ['Contoh: Kelas 10 A'],
        'level': ['SMA'], # Pilihan: SD, SMP, SMA
    }
    df = pd.DataFrame(data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Data Kelas')
    
    output.seek(0)
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="template_import_kelas.xlsx"'
    return response

@login_required
def class_import_view(request):
    """
    Mengimpor data kelas baru dari file Excel.
    """
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        
        try:
            df = pd.read_excel(excel_file)
            
            if 'name' not in df.columns or 'level' not in df.columns:
                messages.error(request, "File Excel harus memiliki kolom 'name' dan 'level'.")
                return redirect('class_management')

            created_count = 0
            warnings = []
            
            for index, row in df.iterrows():
                name = str(row['name']).strip()
                level = str(row['level']).strip().upper()

                if not name or not level:
                    warnings.append(f"Baris {index + 2}: Data nama atau level kosong.")
                    continue

                if level not in ['SD', 'SMP', 'SMA']:
                    warnings.append(f"Baris {index + 2}: Jenjang '{level}' tidak valid. Gunakan SD, SMP, atau SMA.")
                    continue

                # Buat kelas baru jika belum ada, atau lewati jika sudah ada
                _, created = ClassGroup.objects.get_or_create(
                    name=name,
                    defaults={'level': level}
                )
                if created:
                    created_count += 1
            
            if created_count > 0:
                messages.success(request, f'{created_count} kelas baru berhasil dibuat dari file Excel.')
            if warnings:
                messages.warning(request, "Beberapa data tidak dapat diproses:<br>" + "<br>".join(warnings))
            if created_count == 0 and not warnings:
                messages.info(request, "Tidak ada kelas baru yang dibuat (kemungkinan semua sudah ada).")

        except Exception as e:
            messages.error(request, f'Terjadi error saat memproses file: {e}')

    return redirect('class_management')

def class_update_view(request, pk):
    """
    View untuk mengedit data kelas.
    """
    class_instance = get_object_or_404(ClassGroup, pk=pk)
    if request.method == 'POST':
        form = ClassGroupForm(request.POST, instance=class_instance)
        if form.is_valid():
            form.save()
            messages.success(request, f'Data kelas {class_instance.name} berhasil diperbarui!')
            return redirect('class_management')
    else:
        form = ClassGroupForm(instance=class_instance)

    context = {
        'form': form,
        'page_title': 'Edit Kelas'
    }
    # Kita bisa gunakan ulang template user_form.html atau buat yang baru
    return render(request, 'dashboard/class_form.html', context)


@login_required
@require_POST # Memastikan fungsi ini hanya bisa diakses via metode POST untuk keamanan
def class_delete_view(request, pk):
    """
    Menangani logika untuk menghapus sebuah kelas.
    """
    # Ambil kelas yang akan dihapus, atau tampilkan 404 jika tidak ada
    class_to_delete = get_object_or_404(ClassGroup, pk=pk)
    
    class_name = class_to_delete.name
    class_to_delete.delete() # Hapus objek dari database
    
    messages.success(request, f"Kelas '{class_name}' telah berhasil dihapus.")
    
    # Arahkan kembali ke halaman manajemen kelas
    return redirect('class_management')


def export_users_to_excel(request):
    """
    View untuk men-download template Excel kosong untuk data pengguna.
    """
    # PERBAIKAN: Buat DataFrame kosong terlebih dahulu
    columns = {
        'username': [],
        'first_name': [],
        'last_name': [],
        'email': [],
        'password': [],
        'role': [],
    }
    df = pd.DataFrame(columns)
    
    # PERBAIKAN: Ganti nama kolom 'role' untuk memberikan petunjuk
    df.rename(columns={'role': 'role (Pilihan: SISWA, GURU, WALIKELAS, MANAGER, ADMIN)'}, inplace=True)


    # Buat file Excel di dalam memori
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Users')
    
    # Siapkan response untuk di-download oleh browser
    output.seek(0)
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="template_pengguna.xlsx"'
    return response


def import_users_from_excel(request):
    """
    View untuk meng-handle upload file Excel dan mengimpor data pengguna.
    """
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']

        if not excel_file.name.endswith(('.xls', '.xlsx')):
            messages.error(request, 'Format file tidak valid. Harap unggah file Excel (.xls atau .xlsx).')
            return redirect('user_management')

        try:
            df = pd.read_excel(excel_file)
            
            # Mengganti nama kolom petunjuk kembali ke 'role' jika ada
            if 'role (Pilihan: SISWA, GURU, WALIKELAS, MANAGER, ADMIN)' in df.columns:
                df.rename(columns={'role (Pilihan: SISWA, GURU, WALIKELAS, MANAGER, ADMIN)': 'role'}, inplace=True)

            for index, row in df.iterrows():
                user = User(
                    username=row['username'],
                    first_name=row['first_name'],
                    last_name=row['last_name'],
                    email=row['email'],
                    role=str(row['role']).upper() # Pastikan role dalam format uppercase
                )
                user.set_password(str(row['password']))
                user.save()

            messages.success(request, 'Data pengguna berhasil diimpor dari Excel.')

        except Exception as e:
            messages.error(request, f'Terjadi error saat mengimpor: {e}')

    return redirect('user_management')


@login_required
def import_MataPelajarans_view(request):
    """
    View untuk menangani upload dan impor data mata pelajaran dari file Excel.
    """
    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "Tidak ada file yang diunggah.")
            # Ganti 'MataPelajaran_management' dengan nama URL halaman manajemen mapel Anda
            return redirect('MataPelajaran_management') 

        try:
            workbook = openpyxl.load_workbook(excel_file)
            worksheet = workbook.active
            
            # Lewati baris header
            for i, row in enumerate(worksheet.iter_rows(min_row=2, values_only=True)):
                # Pastikan urutan kolom di Excel sama: Nama, Level, Deskripsi
                name, level, description = row
                
                # Gunakan update_or_create untuk menghindari duplikat
                # Ini akan membuat mapel baru jika belum ada, atau memperbarui jika sudah ada
                obj, created = MataPelajaran.objects.update_or_create(
                    nama=name,
                    defaults={
                        # Sesuaikan nama field ini dengan model MataPelajaran Anda
                        'level': level,
                        'description': description
                    }
                )
                if created:
                    messages.success(request, f"Mata pelajaran {name} berhasil ditambahkan.")
                else:
                    messages.info(request, f"Mata pelajaran {name} berhasil diperbarui.")

        except Exception as e:
            messages.error(request, f"Terjadi kesalahan saat mengimpor file: {e}")
        
        return redirect('MataPelajaran_management') # Ganti dengan nama URL Anda
    
    return redirect('MataPelajaran_management') # Ganti dengan nama URL Anda

def MataPelajaran_update_view(request, pk):
    """
    View untuk mengedit data mata pelajaran.
    """
    MataPelajaran = get_object_or_404(MataPelajaran, pk=pk)
    if request.method == 'POST':
        form = MataPelajaranForm(request.POST, instance=MataPelajaran)
        if form.is_valid():
            form.save()
            messages.success(request, f'Mata pelajaran {MataPelajaran.name} berhasil diperbarui!')
            return redirect('MataPelajaran_management')
    else:
        form = MataPelajaranForm(instance=MataPelajaran)

    context = {
        'form': form,
        'page_title': 'Edit Mata Pelajaran'
    }
    return render(request, 'dashboard/MataPelajaran_form.html', context)


def MataPelajaran_delete_view(request, pk):
    """
    View untuk menghapus mata pelajaran.
    """
    MataPelajaran = get_object_or_404(MataPelajaran, pk=pk)
    if request.method == 'POST':
        MataPelajaran_name = MataPelajaran.name
        MataPelajaran.delete()
        messages.success(request, f'Mata pelajaran {MataPelajaran_name} berhasil dihapus.')
        return redirect('MataPelajaran_management')

    context = {
        'MataPelajaran_to_delete': MataPelajaran
    }
    return render(request, 'dashboard/MataPelajaran_confirm_delete.html', context)

@login_required
def question_management_view(request):
    """
    Menampilkan form tambah soal dan daftar soal yang sudah dikelompokkan
    berdasarkan mata pelajaran.
    """
    if request.method == 'POST':
        form = QuestionForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Soal baru berhasil ditambahkan!')
            return redirect('question_management') 
    else:
        form = QuestionForm()

    # Ambil semua soal dan urutkan berdasarkan nama mapel untuk pengelompokan
    all_questions = Question.objects.select_related('matapelajaran').order_by('matapelajaran__nama')

    # Kelompokkan soal ke dalam dictionary
    grouped_questions = defaultdict(list)
    for question in all_questions:
        grouped_questions[question.matapelajaran].append(question)
    
    # Ubah kembali ke dictionary biasa untuk dikirim ke template
    sorted_grouped_questions = dict(sorted(grouped_questions.items(), key=lambda item: item[0].nama))

    context = {
        'form': form,
        'grouped_questions': sorted_grouped_questions, # Kirim data yang sudah dikelompokkan
    }
    return render(request, 'dashboard/question_management.html', context)

@login_required
# Di dalam dashboard/views.py
def question_update_view(request, pk):
    question = get_object_or_404(Question, pk=pk)
    
    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        
        # --- DEBUGGING TAHAP 1: APAKAH FILE DIKIRIM OLEH BROWSER? ---
        print("\n" + "="*20 + " DEBUGGING UPLOAD DIMULAI " + "="*20)
        if request.FILES:
            print(f"✅ request.FILES berisi: {request.FILES}")
        else:
            print("❌ request.FILES KOSONG. Cek 'enctype' di tag <form> HTML Anda.")
        print("="*65 + "\n")
        # --- AKHIR TAHAP 1 ---

        if form.is_valid():
            # --- DEBUGGING TAHAP 2: APAKAH FILE VALID MENURUT FORM? ---
            print("\n" + "="*20 + " FORM IS VALID " + "="*20)
            image_file = form.cleaned_data.get('question_image', None)
            if image_file:
                print(f"✅ Gambar '{image_file.name}' ditemukan di form.cleaned_data.")
            else:
                print("⚠️ Tidak ada gambar baru yang di-upload atau gambar dihapus.")
            print("="*65 + "\n")
            # --- AKHIR TAHAP 2 ---

            try:
                # --- DEBUGGING TAHAP 3: APAKAH PROSES SIMPAN BERHASIL? ---
                print("Mencoba menjalankan form.save()...")
                form.save()
                print("✅ form.save() BERHASIL dijalankan.")
                # --- AKHIR TAHAP 3 ---

                messages.success(request, 'Soal berhasil diperbarui.')
                
                # =================================================================
                # PERBAIKAN: Cek apakah soal memiliki relasi ke ujian sebelum redirect
                # untuk menghindari AttributeError.
                # =================================================================
                if hasattr(question, 'exam') and question.exam:
                    return redirect('exam_detail', pk=question.exam.pk)
                else:
                    # Jika tidak ada ujian, redirect ke halaman manajemen soal utama.
                    return redirect('question_management')

            except Exception as e:
                # --- DEBUGGING TAHAP 4: JIKA ADA ERROR SAAT MENYIMPAN ---
                print("\n" + "="*20 + " !!! ERROR SAAT form.save() !!! " + "="*20)
                print(f"❌ Terjadi error: {e}")
                print("="*65 + "\n")
                # --- AKHIR TAHAP 4 ---

        else:
            # --- DEBUGGING TAHAP 5: JIKA FORM TIDAK VALID ---
            print("\n" + "="*20 + " FORM TIDAK VALID " + "="*20)
            print("❌ Error pada form:")
            print(form.errors.as_json())
            print("="*65 + "\n")
            # --- AKHIR TAHAP 5 ---
            
    else:
        form = QuestionForm(instance=question)
    
    context = {
        'form': form,
        'page_title': f"Edit Soal #{question.pk}",
        'exam': question.exam if hasattr(question, 'exam') else None
    }
    return render(request, 'dashboard/question_form.html', context)


@login_required
def question_edit_view(request, pk):
    question = get_object_or_404(Question, pk=pk)
    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            form.save()
            messages.success(request, 'Soal berhasil diperbarui.')
            return redirect('exam_detail', pk=question.exam.pk)
        else:
            # Tambahkan pesan error jika form tidak valid
            messages.error(request, 'Terjadi kesalahan. Periksa kembali isian Anda.')
    else:
        form = QuestionForm(instance=question)
    
    context = {
        'form': form,
        'question': question
    }
    # Ganti nama template yang di-render
    return render(request, 'dashboard/question_edit_modern.html', context)

def question_delete_view(request, pk):
    """
    Menghapus soal.
    """
    question = get_object_or_404(Question, pk=pk)
    if request.method == 'POST':
        question.delete()
        messages.success(request, 'Soal berhasil dihapus.')
        return redirect('question_management')

    context = {
        'question_to_delete': question
    }
    return render(request, 'dashboard/question_confirm_delete.html', context)

def export_questions_word_template(request):
    """
    Menyediakan file template Word statis untuk diunduh.
    """
    file_path = os.path.join(settings.BASE_DIR, 'static', 'templates', 'template_soal.docx')
    try:
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='template_import_soal.docx')
    except FileNotFoundError:
        messages.error(request, "File template tidak ditemukan. Harap hubungi administrator.")
        return redirect('question_management')

@transaction.atomic # Menambahkan decorator transaksi
def import_questions_from_word(request, exam_pk):
    """
    Mengimpor soal dari file Word.
    Versi ini mendukung:
    - Ekstraksi gambar DAN RUMUS (sebagai gambar) dari sel pertanyaan.
    - Soal berbasis bacaan (stimulus).
    - Berbagai tipe soal sesuai template.
    """
    exam = get_object_or_404(Exam, pk=exam_pk)
    
    if request.method != 'POST' or not request.FILES.get('word_file'):
        return redirect('exam_detail', pk=exam_pk)

    word_file = request.FILES['word_file']
    if not word_file.name.endswith('.docx'):
        messages.error(request, 'Format file tidak valid. Harap unggah file Word (.docx).')
        return redirect('exam_detail', pk=exam_pk)

    try:
        document = docx.Document(word_file)
        if not document.tables:
            messages.error(request, "File Word tidak berisi tabel format yang benar.")
            return redirect('exam_detail', pk=exam_pk)

        table = document.tables[0]
        created_count = 0
        
        # Menggunakan transaction.atomic di sini memastikan semua soal berhasil atau tidak sama sekali
        for i, row in enumerate(table.rows[1:]): # Lewati header
            try:
                # Lewati baris yang sepenuhnya kosong
                if not any(cell.text.strip() for cell in row.cells):
                    continue

                # --- EKSTRAKSI DATA DARI SEL ---
                level = row.cells[0].text.strip().upper() or 'SEDANG'
                bobot_text = row.cells[1].text.strip()
                tipe_soal = row.cells[2].text.strip().upper().replace(" ", "_")
                question_cell = row.cells[3]
                jawaban_text = row.cells[4].text.strip()
                bobot = int(bobot_text) if bobot_text.isdigit() else 10
                
                # --- MEMBANGUN TEKS PERTANYAAN SEBAGAI HTML ---
                question_html_content = ""
                for paragraph in question_cell.paragraphs:
                    p_html = "<p>"
                    for run in paragraph.runs:
                        # Cek apakah 'run' berisi gambar
                        if 'r:embed' in run._r.xml or 'v:imagedata' in run._r.xml:
                            try:
                                embed_id = run._r.xpath('.//a:blip/@r:embed')[0]
                                image_part = document.part.related_parts[embed_id]
                                image_bytes = image_part.blob
                                
                                image_ext = image_part.content_type.split('/')[-1]
                                image_filename = f"question_images/{uuid.uuid4()}.{image_ext}"
                                
                                # PENINGKATAN: Simpan gambar langsung menggunakan storage Django
                                file_path = default_storage.save(image_filename, ContentFile(image_bytes))
                                image_url = default_storage.url(file_path)
                                
                                p_html += f'<img src="{image_url}" alt="Gambar Soal" style="max-width:90%; height:auto;" />'

                            except (IndexError, KeyError):
                                continue # Lanjutkan jika gambar tidak bisa diproses
                        else:
                            p_html += html.escape(run.text)
                    p_html += "</p>"
                    question_html_content += p_html

                options_data = "{}"
                correct_answer_data = jawaban_text

                # --- LOGIKA PEMROSESAN BERDASARKAN TIPE SOAL ---
                #--- PERBAIKAN: Logika parsing opsi sesuai template baru ---
                if tipe_soal == 'MENJODOHKAN':
                    prompts_list = []
                    # Kolom A, B, C (index 5, 6, 7) untuk sisi kiri
                    for cell in row.cells[5:8]:
                        item = cell.text.strip()
                        if ':' in item:
                            parts = item.split(':', 1)
                            prompts_list.append({'id': parts[0].strip(), 'text': parts[1].strip()})
                    
                    choices_list = []
                    # Kolom D, E, F (index 8, 9, 10) untuk sisi kanan
                    for cell in row.cells[8:11]:
                        item = cell.text.strip()
                        if ':' in item:
                            parts = item.split(':', 1)
                            choices_list.append({'id': parts[0].strip(), 'text': parts[1].strip()})

                    if not prompts_list or not choices_list: continue
                        
                    options_data = json.dumps({'prompts': prompts_list, 'choices': choices_list})
                    answer_dict = { pair.split('-', 1)[0].strip(): pair.split('-', 1)[1].strip() for pair in jawaban_text.split(';') if '-' in pair }
                    correct_answer_data = json.dumps(answer_dict)

                elif tipe_soal in ['PILIHAN_GANDA', 'BENAR_SALAH']:
                    options_dict = {}
                    opsi_labels = ['A', 'B', 'C', 'D', 'E']
                    # Kolom A, B, C, D, E (index 5, 6, 7, 8, 9)
                    for j, label in enumerate(opsi_labels):
                        if (5 + j) < len(row.cells):
                            opsi_text = row.cells[5 + j].text.strip()
                            if opsi_text:
                                options_dict[label] = opsi_text
                    
                    options_data = json.dumps(options_dict)
                    correct_answer_data = jawaban_text.upper()
                
                # Untuk tipe soal ESAI dan ISIAN_SINGKAT, biarkan kolom opsi kosong

                # Buat objek Question di database
                Question.objects.create(
                    exam=exam,
                    matapelajaran=exam.MataPelajaran, # Asumsi 'MataPelajaran' adalah nama atribut yang benar
                    question_text=question_html_content,
                    question_type=tipe_soal,
                    difficulty=level,
                    weight=bobot,
                    options=options_data,
                    correct_answer=correct_answer_data
                )
                created_count += 1

            except IndexError:
                messages.warning(request, f"Gagal memproses baris {i + 2}: Jumlah kolom tidak sesuai format.")
                # Karena ada @transaction.atomic, error di sini akan membatalkan semua impor dari file ini
                raise Exception(f"Error di baris {i+2}, impor dibatalkan.")
            except Exception as e:
                messages.warning(request, f"Gagal memproses baris {i + 2} karena error: {e}")
                raise Exception(f"Error di baris {i+2}, impor dibatalkan.")
        
        if created_count > 0:
            messages.success(request, f'{created_count} soal berhasil diimpor.')
        else:
            messages.warning(request, 'Tidak ada soal baru yang berhasil diimpor. Periksa format file Anda.')

    except Exception as e:
        messages.error(request, f'Terjadi error saat membaca file: {e}')

    return redirect('exam_detail', pk=exam_pk)

@login_required
def exam_management_view(request):
    # Bagian 1: Logika untuk menangani pembuatan ujian baru (Method POST)
    if request.method == 'POST':
        form = ExamForm(request.POST)
        if form.is_valid():
            exam = form.save(commit=False)
            exam.created_by = request.user
            exam.save()
            messages.success(request, 'Ujian baru berhasil dibuat.')
            return redirect('exam_management')
    else:
        form = ExamForm()

    # Bagian 2: Logika untuk menampilkan daftar ujian (Method GET)
    user = request.user
    now = timezone.now()
    all_exams = Exam.objects.none() # Mulai dengan queryset kosong

    # Logika spesifik untuk SISWA
    if user.role == 'SISWA':
        if hasattr(user, 'class_group') and user.class_group:
            all_exams = Exam.objects.filter(
                class_group=user.class_group,
                start_time__lte=now,
            ).annotate(
                has_completed=Count('examattempt', filter=Q(examattempt__student=user, examattempt__is_completed=True))
            ).select_related('MataPelajaran', 'class_group').order_by('-start_time')
    
    # Logika spesifik untuk GURU / WALI KELAS (yang bukan superadmin)
    elif not user.is_superuser:
        all_exams = Exam.objects.filter(created_by=user).annotate(
            total_students_in_class_count=Count('class_group__students', filter=Q(class_group__students__role='SISWA')),
            completed_attempts_count=Count('examattempt', filter=Q(examattempt__is_completed=True))
        ).select_related('MataPelajaran', 'class_group', 'created_by').order_by('-start_time')

    # Logika untuk ADMIN / SUPERUSER
    else:
        all_exams = Exam.objects.all().annotate(
            total_students_in_class_count=Count('class_group__students', filter=Q(class_group__students__role='SISWA')),
            completed_attempts_count=Count('examattempt', filter=Q(examattempt__is_completed=True))
        ).select_related('MataPelajaran', 'class_group', 'created_by').order_by('-start_time')

    # Bagian 3: Mengelompokkan ujian berdasarkan jenjang (berlaku untuk semua peran)
    grouped_exams = defaultdict(list)
    for exam in all_exams:
        if exam.class_group and hasattr(exam.class_group, 'level'):
            grouped_exams[exam.class_group.get_level_display()].append(exam)
        else:
            grouped_exams["Tanpa Jenjang"].append(exam)

    context = {
        'form': form,
        'grouped_exams': dict(sorted(grouped_exams.items())),
        'page_title': 'Manajemen Ujian',
    }
    return render(request, 'dashboard/exam_management.html', context)

@login_required
def exam_detail_view(request, pk):
    exam = get_object_or_404(Exam, pk=pk)

    # =======================================================
    # BAGIAN BARU: Logika Pembuatan Token (Method POST)
    # =======================================================
    if request.method == 'POST':
        # Cek apakah aksi yang diminta adalah 'generate_tokens'
        if 'generate_tokens' in request.POST:
            try:
                # Ambil jumlah token yang ingin dibuat dari form
                num_tokens_str = request.POST.get('num_tokens', '1')
                num_tokens = int(num_tokens_str)

                if num_tokens <= 0 or num_tokens > 100: # Batasi pembuatan maksimal 100 sekali generate
                     messages.error(request, 'Jumlah token harus antara 1 dan 100.')
                else:
                    # Buat token sebanyak jumlah yang diminta
                    for _ in range(num_tokens):
                        # Generate string acak yang aman dan unik
                        new_token_str = secrets.token_hex(4).upper()
                        # Pastikan token belum ada untuk ujian ini (sangat kecil kemungkinannya, tapi best practice)
                        while Token.objects.filter(exam=exam, token=new_token_str).exists():
                            new_token_str = secrets.token_hex(4).upper()
                        
                        # Buat objek token baru dan simpan
                        Token.objects.create(exam=exam, token=new_token_str)
                    
                    messages.success(request, f'{num_tokens} token baru berhasil dibuat.')

            except (ValueError, TypeError):
                messages.error(request, 'Jumlah token tidak valid.')
            
            # Arahkan kembali ke halaman yang sama untuk me-refresh halaman
            return redirect('exam_detail', pk=exam.pk)

    # =======================================================
    # BAGIAN BARU: Ambil Daftar Token yang Sudah Ada
    # =======================================================
    tokens = exam.token_set.all().order_by('-created_at') # Ambil semua token untuk exam ini

    # --- Kode lama Anda untuk mengambil soal (tidak perlu diubah) ---
    questions = exam.questions.order_by('id')
    questions_json = json.dumps(
        [{
            'id': q.id,
            'text': q.question_text,
            'type': q.get_question_type_display(),
            'options': q.options,
            'correct_answer': q.correct_answer,
            'difficulty': q.get_difficulty_display(),
            'weight': q.weight,
            'edit_url': reverse('question_edit', args=[q.pk]),
        } for q in questions],
        cls=DjangoJSONEncoder
    )

    context = {
        'exam': exam,
        'questions': questions,
        'questions_json': questions_json,
        'tokens': tokens, # <-- KIRIM TOKEN KE TEMPLATE
    }
    return render(request, 'dashboard/exam_detail.html', context)


@login_required
def exam_update_view(request, pk):
    # 1. Ambil objek ujian yang akan diedit
    exam = get_object_or_404(Exam, pk=pk)
    
    if request.method == 'POST':
        form = ExamForm(request.POST, instance=exam)
        if form.is_valid():
            form.save()
            messages.success(request, 'Ujian berhasil diperbarui.')
            return redirect('exam_management') # Arahkan kembali ke daftar ujian
    else:
        # Isi form dengan data dari objek ujian yang ada
        form = ExamForm(instance=exam)

    context = {
        'form': form,
        'exam': exam, # 2. KIRIM OBJEK EXAM KE TEMPLATE (INI BAGIAN PENTING)
        'page_title': f'Edit Ujian: {exam.name}'
    }
    
    # Gunakan template form yang sesuai, misalnya 'exam_form.html'
    return render(request, 'dashboard/exam_form.html', context)

def exam_delete_view(request, pk):
    """
    View untuk menghapus sebuah ujian.
    """
    exam = get_object_or_404(Exam, pk=pk)
    if request.method == 'POST':
        exam_name = exam.name
        exam.delete()
        messages.success(request, f'Ujian "{exam_name}" berhasil dihapus.')
        return redirect('exam_management')

    context = {
        'exam_to_delete': exam
    }
    # Kita akan buat template baru untuk ini
    return render(request, 'dashboard/exam_confirm_delete.html', context)

@transaction.atomic
@login_required
def take_exam_view(request, pk):
    exam = get_object_or_404(Exam, pk=pk)
    questions = exam.questions.all()

    questions_as_dicts = []
    for q in questions:
        options_list = []

        if q.question_type == "PILIHAN_GANDA":
            # Jika options adalah JSONField (dict)
            if isinstance(q.options, dict):
                for key, val in q.options.items():
                    options_list.append({"pk": key, "text": val})
            # Jika options adalah ManyToManyField
            else:
                for opt in q.options.all():
                    options_list.append({"pk": opt.pk, "text": opt.option_text})

        elif q.question_type == "MENJODOHKAN":
            # contoh: {"left": ["Kucing","Burung"], "right": ["Mamalia","Unggas"]}
            options_list = q.options  

        questions_as_dicts.append({
            "id": q.id,
            "text": q.question_text,
            "type": q.question_type,
            "options": options_list,
            "image_url": q.question_image.url if q.question_image else None
        })

    questions_json = json.dumps(questions_as_dicts, cls=DjangoJSONEncoder)

    context = {
        "exam": exam,
        "questions_json": questions_json,
    }
    return render(request, "dashboard/student/take_exam.html", context)

@login_required
@require_POST # View ini hanya menerima request POST
def exam_submit_view(request, pk): # NAMA FUNGSI DIPERBAIKI
    """
    Menangani pengiriman jawaban ujian dari siswa (menerima data JSON).
    Menggabungkan logika perhitungan skor dari kode Anda.
    """
    try:
        exam = get_object_or_404(Exam, pk=pk)
        # Dapatkan atau buat objek pengerjaan ujian
        attempt, created = ExamAttempt.objects.get_or_create(student=request.user, exam=exam)
        
        # Ambil data jawaban yang dikirim oleh JavaScript dari request body
        answers_data = json.loads(request.body).get('answers', {})

        # Hapus jawaban lama jika siswa submit ulang untuk menghindari duplikasi
        Answer.objects.filter(attempt=attempt).delete()

        needs_manual_correction = False
        correct_count = 0
        total_graded_questions = 0

        # Iterasi melalui jawaban yang diterima dari JSON
        for question_pk, answer_value in answers_data.items():
            question = get_object_or_404(Question, pk=int(question_pk))
            
            is_correct = None
            
            # Logika untuk tipe soal yang bisa dinilai otomatis
            if question.question_type in ['PILIHAN_GANDA', 'BENAR_SALAH']:
                is_correct = (str(answer_value) == str(question.correct_answer))
                if is_correct:
                    correct_count += 1
                total_graded_questions += 1
            
            # Logika untuk tipe soal yang butuh koreksi manual
            elif question.question_type == 'ESSAY':
                needs_manual_correction = True

            # Simpan setiap jawaban ke database
            Answer.objects.create(
                attempt=attempt,
                question=question,
                # 'answer_value' bisa berisi pilihan (misal: 'A') atau teks esai
                answer_text=str(answer_value), 
                is_correct=is_correct
            )
        
        # Update skor dan status pengerjaan
        if not needs_manual_correction and total_graded_questions > 0:
            attempt.score = (correct_count / total_graded_questions) * 100
        else:
            # Atur skor menjadi None jika ada esai yang perlu dikoreksi manual
            attempt.score = None
            
        attempt.is_completed = True
        attempt.end_time = timezone.now()
        attempt.save()

        # Kirim respons sukses kembali ke JavaScript
        return JsonResponse({
            'status': 'success',
            'message': 'Ujian berhasil diselesaikan!',
            'redirect_url': reverse('dashboard_siswa') # Mengarahkan ke dashboard siswa sesuai kode Anda
        })

    except ExamAttempt.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Sesi ujian tidak ditemukan.'}, status=404)
    except Exception as e:
        # Catat error untuk debugging
        print(f"Error saat submit ujian: {e}")
        return JsonResponse({'status': 'error', 'message': 'Terjadi kesalahan internal.'}, status=500)
    
def home_view(request):
    """
    Menampilkan halaman utama jika pengguna belum login.
    Jika sudah login, arahkan langsung ke dasbor.
    """
    if request.user.is_authenticated:
        return redirect('user_management') # Arahkan ke dasbor jika sudah login
    
    # Tampilkan halaman home.html jika belum login
    return render(request, 'dashboard/main/home.html')

def teacher_profile_view(request, pk):
    teacher = get_object_or_404(User, pk=pk, role__in=[User.Role.GURU, User.Role.WALIKELAS])
    assignments = TeacherAssignment.objects.filter(teacher=teacher).select_related('MataPelajaran', 'class_group')
    evaluations = teacher.received_evaluations.select_related('evaluator').all()

    # Inisialisasi kedua form
    evaluation_form = EvaluationForm()
    profile_update_form = TeacherProfileUpdateForm(instance=teacher)

    if request.method == 'POST' and request.user.role == 'MANAGER':
        # Cek tombol mana yang ditekan berdasarkan atribut 'name'
        if 'submit_evaluation' in request.POST:
            evaluation_form = EvaluationForm(request.POST)
            if evaluation_form.is_valid():
                evaluation = evaluation_form.save(commit=False)
                evaluation.evaluator = request.user
                evaluation.teacher = teacher
                evaluation.save()
                messages.success(request, 'Evaluasi berhasil ditambahkan.')
                return redirect('teacher_profile', pk=teacher.pk)
        
        elif 'update_profile' in request.POST:
            profile_update_form = TeacherProfileUpdateForm(request.POST, instance=teacher)
            if profile_update_form.is_valid():
                profile_update_form.save()
                messages.success(request, 'Profil guru berhasil diperbarui.')
                return redirect('teacher_profile', pk=teacher.pk)

    context = {
        'teacher': teacher,
        'assignments': assignments,
        'evaluations': evaluations,
        'evaluation_form': evaluation_form,
        'profile_update_form': profile_update_form,
    }
    return render(request, 'dashboard/teacher/teacher_profile.html', context)
@login_required
def exam_start_page_view(request, attempt_pk):
    attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
    exam = attempt.exam

    # Keamanan: Jika attempt sudah selesai, tendang keluar
    if attempt.is_completed:
        messages.info(request, "Ujian telah selesai.")
        return redirect('exam_result_detail', attempt_pk=attempt.pk) # Arahkan ke hasil

    # Perhitungan sisa waktu
    total_duration_seconds = exam.duration_minutes * 60
    elapsed_seconds = (timezone.now() - attempt.start_time).total_seconds()
    time_left = total_duration_seconds - elapsed_seconds
    if time_left < 0:
        time_left = 0

    # Pengambilan dan pemrosesan soal (logika parsing Anda sudah bagus, kita pakai lagi)
    question_queryset = Question.objects.filter(exam=exam).order_by('id').distinct()
    questions_as_dictionaries = []

    for question in question_queryset:
        options_dict = None
        if isinstance(question.options, dict):
            options_dict = question.options
        elif isinstance(question.options, str) and question.options.strip():
            try:
                options_dict = json.loads(question.options)
            except json.JSONDecodeError:
                try:
                    options_dict = ast.literal_eval(question.options)
                except (ValueError, SyntaxError):
                    options_dict = None

        parsed_options = None # Default value
        if question.question_type == 'PILIHAN_GANDA':
            parsed_options = []
            if isinstance(options_dict, dict):
                for key, value in options_dict.items():
                    parsed_options.append({'id': str(key), 'text': value})
        
        elif question.question_type == 'MENJODOHKAN':
            # Logika parsing menjodohkan Anda...
            # (Saya persingkat untuk kejelasan, logika Anda sudah benar)
            parsed_options = {'prompts': [], 'choices': []}
            if isinstance(options_dict, dict):
                 # ... (logika parsing menjodohkan Anda)
                 parsed_options = options_dict # Asumsi format sudah benar

        elif question.question_type == 'BENAR_SALAH':
            parsed_options = [
                {'id': 'Benar', 'text': 'Benar'},
                {'id': 'Salah', 'text': 'Salah'}
            ]

        questions_as_dictionaries.append({
            'pk': question.pk,
            'text': question.question_text,
            'type': question.question_type,
            'image_url': question.question_image.url if question.question_image else None,
            'audio_url': question.question_audio.url if question.question_audio else None,
            'options': parsed_options
        })

    context = {
        'exam': exam,
        'attempt': attempt, # Kirim objek attempt
        'time_left': time_left,
        'questions_json': json.dumps(questions_as_dictionaries),
        'saved_answers_json': json.dumps(attempt.current_answers or {})
    }
    return render(request, 'dashboard/exam_start.html', context)

@login_required
def exam_token_gate_view(request, pk):
    exam = get_object_or_404(Exam, pk=pk)
    context = {'exam': exam}

    # Keamanan: Cek waktu ujian & status selesai
    if ExamAttempt.objects.filter(student=request.user, exam=exam, is_completed=True).exists():
        messages.warning(request, 'Anda sudah menyelesaikan ujian ini.')
        return redirect('exam_management') # Ganti dengan URL name yang sesuai
    now = timezone.now()
    if now > exam.end_time:
        messages.error(request, "Waktu ujian telah berakhir.")
        return redirect('exam_management') # Ganti dengan URL name yang sesuai
    if now < exam.start_time:
        messages.error(request, "Ujian ini belum dimulai.")
        return redirect('exam_management') # Ganti dengan URL name yang sesuai

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'validate_token':
            token_str = request.POST.get('token', '').strip().upper()
            if not token_str:
                messages.error(request, 'Token tidak boleh kosong.')
                return redirect('exam_token_gate', pk=exam.pk)

            # Cari objek token yang valid.
            valid_token_obj = Token.objects.filter(exam=exam, token=token_str).first()

            if not valid_token_obj:
                messages.error(request, 'Token yang Anda masukkan salah.')
                return redirect('exam_token_gate', pk=exam.pk)

            # =======================================================
            # PERUBAHAN UTAMA: Cek apakah token ini sudah pernah dipakai siswa.
            # =======================================================
            if ExamAttempt.objects.filter(student=request.user, exam=exam, token_used=valid_token_obj).exists():
                messages.error(request, "Token ini sudah pernah Anda gunakan untuk memulai ujian dan tidak dapat dipakai lagi.")
                return redirect('exam_token_gate', pk=exam.pk)
            
            # Jika lolos, token valid dan belum pernah dipakai siswa ini.
            # Simpan di session dan tampilkan halaman persiapan.
            request.session[f'exam_{pk}_token_validated'] = token_str
            context['show_prepare_page'] = True
            return render(request, 'dashboard/exam_token_gate.html', context)

        # --- LANGKAH 2: Mulai Ujian ---
        elif action == 'start_exam':
            token_str = request.session.get(f'exam_{pk}_token_validated')
            if not token_str:
                messages.error(request, 'Sesi token tidak valid. Silakan ulangi.')
                return redirect('exam_token_gate', pk=exam.pk)

            valid_token_obj = Token.objects.filter(exam=exam, token=token_str).first()
            
            # Buat percobaan ujian (attempt) yang baru.
            # Kita menggunakan create() karena sudah yakin ini adalah percobaan baru.
            attempt = ExamAttempt.objects.create(
                student=request.user,
                exam=exam,
                start_time=timezone.now(),
                token_used=valid_token_obj
            )

            # Hapus session dan arahkan ke halaman soal
            del request.session[f'exam_{pk}_token_validated']
            return redirect('exam_start_page', attempt_pk=attempt.pk)

    # Tampilan Halaman Gerbang Token (Method GET)
    attempt = ExamAttempt.objects.filter(student=request.user, exam=exam, is_completed=False).first()
    if attempt:
        context['page_heading'] = "Ujian Anda Dijeda"
        context['page_subheading'] = "Sistem mendeteksi Anda meninggalkan halaman. Masukkan token yang sama untuk melanjutkan."
        context['button_text'] = "Lanjutkan Ujian"
    else:
        context['page_heading'] = "Anda akan memulai ujian:"
        context['page_subheading'] = "Silakan masukkan token yang diberikan oleh guru Anda untuk memulai."
        context['button_text'] = "Mulai Kerjakan"
        
    return render(request, 'dashboard/exam_token_gate.html', context)
                  
@login_required
def enroll_students_view(request, pk):
    """
    View yang diperbaiki untuk menampilkan daftar siswa di dalam kelas.
    """
    class_group = get_object_or_404(ClassGroup, pk=pk)

    # Logika untuk mendaftarkan siswa dari form manual
    if request.method == 'POST':
        student_pks = request.POST.getlist('students')
        if student_pks:
            students_to_enroll = User.objects.filter(pk__in=student_pks, role='SISWA')
            enrolled_count = students_to_enroll.update(class_group=class_group)
            messages.success(request, f"{enrolled_count} siswa berhasil didaftarkan ke kelas {class_group.name}.")
        else:
            messages.warning(request, "Tidak ada siswa yang dipilih.")
        return redirect('enroll_students', pk=pk)

    # --- PERBAIKAN UTAMA: Mengambil data siswa dengan benar ---
    enrolled_students = User.objects.filter(class_group=class_group, role='SISWA').order_by('first_name')
    unenrolled_students = User.objects.filter(class_group__isnull=True, role='SISWA').order_by('first_name')

    context = {
        'class_group': class_group,
        'enrolled_students': enrolled_students,
        'unenrolled_students': unenrolled_students,
    }
    return render(request, 'dashboard/enroll_students.html', context)


@login_required
def profile_view(request):
    user = request.user

    assignments = None
    evaluations = None
    if hasattr(user, 'role') and user.role in [User.Role.GURU, User.Role.WALIKELAS]:
        assignments = TeacherAssignment.objects.filter(
            teacher=user
        ).select_related('MataPelajaran', 'class_group')
        evaluations = user.received_evaluations.select_related('evaluator').all()

    # inisialisasi form awal
    profile_form = ProfileUpdateForm(instance=user)
    password_change_form = CustomPasswordChangeForm(user)

    if request.method == 'POST':
        if 'update_profile' in request.POST:
            profile_form = ProfileUpdateForm(request.POST, request.FILES, instance=user)
            if profile_form.is_valid():
                profile_form.save()
                messages.success(request, 'Profil Anda berhasil diperbarui.')
                return redirect('profile')
            else:
                messages.error(request, 'Gagal memperbarui profil. Silakan periksa kembali isian Anda.')

        elif 'change_password' in request.POST:
            password_change_form = CustomPasswordChangeForm(user, request.POST)
            if password_change_form.is_valid():
                saved_user = password_change_form.save()
                update_session_auth_hash(request, saved_user)
                messages.success(request, 'Password Anda berhasil diubah.')
                return redirect('profile')
            else:
                messages.error(request, 'Gagal mengubah password. Silakan periksa kembali isian Anda.')

    return render(request, 'dashboard/profile.html', {
        'profile_form': profile_form,
        'password_change_form': password_change_form,  # penting! sama dengan di template
        'assignments': assignments,
        'evaluations': evaluations,
    })

def export_new_student_template(request):
    """
    Men-download template Excel untuk membuat siswa baru, lengkap dengan NIS dan NISN.
    """
    data = {
        'username': ['budi.santoso'],
        'password': ['password_aman_123'],
        'first_name': ['Budi'],
        'last_name': ['Santoso'],
        'email': ['budi.s@email.com'],
        'nis': ['2025001'], # Kolom baru
        'nisn': ['1234567890'],
        'class_name': ['Kelas 10 A'],
        'theme': ['LIGHT'], 
    }
    df = pd.DataFrame(data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Siswa Baru')
    
    output.seek(0)
    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="template_import_siswa_baru.xlsx"'
    return response

def import_new_students_view(request):
    """
    Mengimpor siswa baru dari file Excel, termasuk NIS dan NISN.
    """
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        # ... (validasi file excel) ...

        try:
            df = pd.read_excel(excel_file)
            created_count = 0
            for index, row in df.iterrows():
                # ... (ambil data lain: username, password, dll.) ...
                nis = str(row.get('nis', '')).strip() # Ambil data NIS
                nisn = str(row.get('nisn', '')).strip() # Ambil data NISN
                class_name = str(row['class_name']).strip()
                
                user, created = User.objects.get_or_create(
                    username=str(row['username']).strip(),
                    defaults={
                        'username' : str (row['username']).strip(),
                        'first_name': str(row['first_name']).strip(),
                        'last_name': str(row['last_name']).strip(),
                        'email': str(row['email']).strip(),
                        'nis': nis if nis else None, # Simpan NIS
                        'nisn': nisn if nisn else None, # Simpan NISN
                        'role': 'SISWA',
                        'theme': str(row.get('theme', 'LIGHT')).strip().upper()
                    }
                )
                if created:
                    user.set_password(str(row['password']).strip())
                    user.save()
                    created_count += 1

                # ... (logika pendaftaran ke kelas tidak berubah) ...
            
            messages.success(request, f'{created_count} siswa baru berhasil dibuat dan didaftarkan.')

        except Exception as e:
            messages.error(request, f'Terjadi error: {e}')

    return redirect('user_management')

@login_required
@transaction.atomic # Memastikan semua proses berhasil atau tidak sama sekali
def import_new_students_to_class(request, class_pk):
    """
    Membuat akun siswa baru dari file Excel dan langsung mendaftarkannya 
    ke kelas yang spesifik.
    """
    class_instance = get_object_or_404(ClassGroup, pk=class_pk)
    
    if request.method != 'POST' or not request.FILES.get('excel_file'):
        messages.error(request, 'Permintaan tidak valid.')
        return redirect('enroll_students', pk=class_pk)

    excel_file = request.FILES['excel_file']
    if not excel_file.name.endswith(('.xlsx', '.xls')):
        messages.error(request, 'Format file tidak valid. Harap unggah file Excel (.xlsx atau .xls).')
        return redirect('enroll_students', pk=class_pk)

    try:
        df = pd.read_excel(excel_file)
        
        required_columns = ['username', 'password', 'nis']
        if not all(col in df.columns for col in required_columns):
            messages.error(request, f"File Excel harus memiliki kolom wajib: {', '.join(required_columns)}.")
            return redirect('enroll_students', pk=class_pk)

        created_count = 0
        warnings = []

        for index, row in df.iterrows():
            username = str(row['username']).strip()
            nis = str(row['nis']).strip()

            if not username or not nis:
                warnings.append(f"Baris {index + 2}: 'username' dan 'nis' tidak boleh kosong.")
                continue

            if User.objects.filter(username=username).exists() or User.objects.filter(nis=nis).exists():
                warnings.append(f"Baris {index + 2}: Username '{username}' atau NIS '{nis}' sudah terdaftar.")
                continue

            try:
                # 1. Buat pengguna dengan parameter standar terlebih dahulu
                user = User.objects.create_user(
                    username=username,
                    password=row['password'],
                    first_name=row.get('first_name', ''),
                    last_name=row.get('last_name', ''),
                    email=row.get('email', '')
                )

                # 2. Setelah pengguna dibuat, isi kolom-kolom tambahan
                user.nis = nis
                user.nisn = row.get('nisn', '')
                user.role = 'SISWA'
                user.class_group = class_instance # Tetapkan kelasnya
                user.save() # Simpan perubahan

                created_count += 1
            except Exception as e:
                warnings.append(f"Baris {index + 2}: Gagal membuat pengguna '{username}'. Error: {e}")

        if created_count > 0:
            messages.success(request, f'{created_count} siswa baru berhasil dibuat dan didaftarkan ke kelas {class_instance.name}.')
        if warnings:
            messages.warning(request, "Beberapa data tidak dapat diproses:<br>" + "<br>".join(warnings))
        if created_count == 0 and not warnings:
            messages.info(request, "Tidak ada siswa baru yang dibuat dari file yang diunggah.")

    except Exception as e:
        messages.error(request, f'Terjadi error saat memproses file: {e}')

    return redirect('enroll_students', pk=class_pk)

@login_required
def correction_list_view(request): # Nama fungsi disesuaikan dengan error Anda
    """
    Menampilkan daftar ujian yang perlu dikoreksi, HANYA untuk ujian
    yang dibuat oleh guru yang sedang login.
    """
    
    current_teacher = request.user
    search_query = request.GET.get('search', '')

    # =========================================================================
    # INI BAGIAN YANG DIPERBAIKI
    # Kita filter berdasarkan 'is_completed' dan 'score' yang masih kosong (NULL)
    # =========================================================================
    base_queryset = ExamAttempt.objects.filter(
        exam__created_by=current_teacher,
        is_completed=True,
        score__isnull=True  # Asumsi: jika skor masih NULL, berarti perlu dikoreksi
    )

    if search_query:
        base_queryset = base_queryset.filter(
            Q(exam__name__icontains=search_query) |
            Q(exam__MataPelajaran__name__icontains=search_query) |
            Q(student__first_name__icontains=search_query) |
            Q(student__last_name__icontains=search_query)
        )

    attempts_to_correct = base_queryset.select_related(
        'student', 'exam', 'exam__MataPelajaran', 'exam__class_group'
    ).order_by('exam__name', 'end_time')

    grouped_attempts = {}
    for attempt in attempts_to_correct:
        if attempt.exam not in grouped_attempts:
            grouped_attempts[attempt.exam] = []
        grouped_attempts[attempt.exam].append(attempt)

    context = {
        'grouped_attempts': grouped_attempts,
        'search_query': search_query,
    }

    return render(request, 'dashboard/teacher/correction_list.html', {
        'grouped_attempts': grouped_attempts,
        'search_query': search_query
    })

def view_evaluations_view(request):
    """
    Halaman khusus untuk guru melihat semua evaluasi yang diterima.
    """
    user = request.user
    # Pastikan hanya guru/wali kelas yang bisa mengakses
    if user.role not in [User.Role.GURU, User.Role.WALIKELAS]:
        # Arahkan ke halaman lain jika bukan guru
        return redirect('profile') 

    evaluations = user.received_evaluations.select_related('evaluator').all()
    context = {
        'evaluations': evaluations,
    }
    return render(request, 'dashboard/teacher/view_evaluations.html', context)

@login_required
@transaction.atomic # Memastikan semua proses berhasil atau tidak sama sekali
def import_new_students_view(request):
    """
    Mengimpor siswa baru dari file Excel dengan pencocokan kelas yang lebih baik.
    """
    redirect_url = request.POST.get('next', 'user_management')

    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        # ... (validasi file excel) ...

        try:
            df = pd.read_excel(excel_file)
            # ... (validasi kolom) ...

            created_count = 0
            for index, row in df.iterrows():
                username = str(row['username']).strip()
                class_name = str(row['class_name']).strip()
                
                # ... (logika pembuatan user tidak berubah) ...
                user, created = User.objects.get_or_create(
                    username=username,
                    defaults={
                        'first_name': str(row['first_name']).strip(),
                        'last_name': str(row['last_name']).strip(),
                        'email': str(row['email']).strip(),
                        'nis': str(row.get('nis', '')).strip(),
                        'nisn': str(row.get('nisn', '')).strip(),
                        'role': 'SISWA'
                    }
                )
                if created:
                    user.set_password(str(row['password']).strip())
                    user.save()
                    created_count += 1

                # =======================================================
                # PERBAIKAN LOGIKA PENCOCOKAN KELAS
                # =======================================================
                try:
                    # Coba cari kelas dengan nama yang sama persis (case-insensitive)
                    class_group = ClassGroup.objects.get(name__iexact=class_name)
                    Enrollment.objects.get_or_create(student=user, class_group=class_group)
                except ClassGroup.DoesNotExist:
                    # Jika tidak ditemukan, beri peringatan
                    messages.warning(request, f"Peringatan di baris {index + 2}: Kelas '{class_name}' tidak ditemukan. Siswa '{username}' dibuat tanpa kelas.")
                # =======================================================
            
            if created_count > 0:
                messages.success(request, f'{created_count} siswa baru berhasil dibuat dan didaftarkan.')
            else:
                messages.info(request, 'Tidak ada siswa baru yang dibuat (kemungkinan semua sudah ada).')

        except Exception as e:
            messages.error(request, f'Terjadi error saat memproses file: {e}')

    # PERBAIKAN: Arahkan kembali ke halaman yang benar
    if 'class' in redirect_url:
        # Ekstrak class_pk dari path seperti '/dashboard/class/35/enroll/'
        try:
            class_pk = redirect_url.split('/')[3]
            return redirect('class_enroll', pk=class_pk)
        except (IndexError, ValueError):
            return redirect('class_management')
    else:
        return redirect(redirect_url)


@login_required
def unenroll_student_view(request, class_pk, student_pk):
    """
    Mengeluarkan (unenroll) seorang siswa dari kelasnya.
    Ini tidak menghapus data siswa, hanya mengatur ulang kelasnya menjadi None.
    """
    # Dapatkan instance kelas dan siswa yang akan dikeluarkan
    class_instance = get_object_or_404(ClassGroup, pk=class_pk)
    student_to_unenroll = get_object_or_404(User, pk=student_pk, role='SISWA')

    # Pastikan siswa tersebut memang ada di kelas ini sebelum dikeluarkan
    if student_to_unenroll.class_group == class_instance:
        # Setel class_group siswa menjadi None (atau null)
        student_to_unenroll.class_group = None
        student_to_unenroll.save(update_fields=['class_group'])
        
        messages.success(request, f"Siswa '{student_to_unenroll.get_full_name()}' telah berhasil dikeluarkan dari kelas {class_instance.name}.")
    else:
        messages.warning(request, f"Siswa '{student_to_unenroll.get_full_name()}' tidak terdaftar di kelas ini.")

    # Redirect kembali ke halaman kelola siswa
    # Pastikan 'enroll_students' adalah nama URL yang benar untuk halaman kelola siswa
    return redirect('enroll_students', pk=class_pk)

@login_required
def user_import_view(request): # Pastikan nama fungsi ini sesuai dengan yang Anda gunakan
    """
    Mengimpor pengguna baru dari file Excel, sekarang dengan logika untuk
    menemukan dan menugaskan kelas berdasarkan 'class_name'.
    """
    # --- PERBAIKAN: Semua kode di dalam fungsi ini sekarang di-indentasi dengan benar ---
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        
        try:
            df = pd.read_excel(excel_file)
            
            required_columns = ['username', 'password', 'nis']
            if not all(col in df.columns for col in required_columns):
                messages.error(request, f"File Excel harus memiliki kolom wajib: {', '.join(required_columns)}.")
                return redirect('user_management')

            created_count = 0
            warnings = []

            for index, row in df.iterrows():
                username = str(row['username']).strip()
                nis = str(row.get('nis', '')).strip()

                if not username:
                    warnings.append(f"Baris {index + 2}: 'username' tidak boleh kosong.")
                    continue

                if User.objects.filter(username=username).exists():
                    warnings.append(f"Baris {index + 2}: Username '{username}' sudah digunakan.")
                    continue
                
                # --- LOGIKA BARU UNTUK KELAS DIMULAI DI SINI ---
                class_instance = None
                class_name = str(row.get('class_name', '')).strip()
                if class_name:
                    try:
                        # Cari kelas di database berdasarkan nama dari Excel
                        class_instance = ClassGroup.objects.get(name__iexact=class_name)
                    except ClassGroup.DoesNotExist:
                        # Jika kelas tidak ditemukan, catat sebagai peringatan
                        warnings.append(f"Baris {index + 2}: Kelas '{class_name}' tidak ditemukan. Siswa '{username}' dibuat tanpa kelas.")
                # --- LOGIKA KELAS SELESAI ---

                try:
                    user = User.objects.create_user(
                        username=username,
                        password=row['password'],
                        first_name=row.get('first_name', ''),
                        last_name=row.get('last_name', ''),
                        email=row.get('email', ''),
                        role=str(row.get('role', 'SISWA')).strip().upper(),
                        nis=nis,
                        nisn=str(row.get('nisn', '')).strip(),
                        # Tugaskan kelas yang sudah ditemukan (atau None jika tidak ketemu)
                        class_group=class_instance 
                    )
                    created_count += 1
                except Exception as e:
                    warnings.append(f"Baris {index + 2}: Gagal membuat pengguna '{username}'. Error: {e}")

            if created_count > 0:
                messages.success(request, f'{created_count} pengguna baru berhasil dibuat.')
            if warnings:
                messages.warning(request, "Beberapa data tidak dapat diproses:<br>" + "<br>".join(warnings))
            if created_count == 0 and not warnings:
                messages.info(request, "Tidak ada pengguna baru yang dibuat.")

        except Exception as e:
            messages.error(request, f'Terjadi error saat memproses file: {e}')

    return redirect('user_management')

@login_required
def export_class_students(request, class_pk):
    """
    Mengekspor daftar siswa dari kelas spesifik ke dalam file Excel.
    """
    class_instance = get_object_or_404(ClassGroup, pk=class_pk)
    
    # Dapatkan semua siswa yang terdaftar di kelas ini, urutkan berdasarkan nama
    students = User.objects.filter(class_group=class_instance, role='SISWA').order_by('first_name')

    if not students.exists():
        messages.warning(request, 'Tidak ada siswa untuk diekspor dari kelas ini.')
        return redirect('enroll_students', pk=class_pk)

    # Siapkan data yang akan dimasukkan ke dalam file Excel
    data = {
        'Nama Lengkap': [s.get_full_name() for s in students],
        'Username': [s.username for s in students],
        'NIS': [s.nis for s in students],
        'NISN': [s.nisn for s in students],
        'Email': [s.email for s in students],
    }
    df = pd.DataFrame(data)

    # Buat file Excel di dalam memori (bukan di hard drive)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=f'Siswa {class_instance.name}')
    
    output.seek(0)

    # Buat response HTTP untuk men-download file Excel
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # Atur nama file yang akan di-download
    response['Content-Disposition'] = f'attachment; filename="daftar_siswa_{class_instance.name}.xlsx"'
    
    return response
@login_required
def user_export(request):
    """
    Membuat dan menyediakan file template Excel kosong untuk impor pengguna baru,
    sesuai dengan format yang diminta.
    """
    # Tentukan kolom-kolom yang dibutuhkan sesuai file Excel Anda
    data = {
        'username': [],
        'password': [],
        'first_name': [],
        'last_name': [],
        'email': [],
        'nis': [],
        'nisn': [],
        'class_name': [], # Kolom baru
        'theme': [],      # Kolom baru
    }
    df = pd.DataFrame(data)

    # Buat file Excel di dalam memori
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Siswa Baru')
        
        # Tambahkan catatan atau instruksi di sheet terpisah
        worksheet_note = writer.book.create_sheet("Instruksi")
        worksheet_note.cell(row=1, column=1, value="PENTING:")
        worksheet_note.cell(row=2, column=1, value="1. Kolom 'username', 'password', dan 'nis' wajib diisi.")
        worksheet_note.cell(row=3, column=1, value="2. Kolom 'class_name' di sini hanya untuk referensi. Siswa akan otomatis dimasukkan ke kelas yang sedang Anda kelola.")
        worksheet_note.cell(row=4, column=1, value="3. Jangan mengubah nama header kolom.")


    output.seek(0)

    # Buat response HTTP untuk men-download file
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # Atur nama file yang akan di-download
    response['Content-Disposition'] = 'attachment; filename="template_import_siswa_baru.xlsx"'
    
    return response
@login_required
def edit_class_view(request, class_id):
    # Dapatkan objek kelas yang akan diedit, atau tampilkan 404 jika tidak ada
    class_group = get_object_or_404(ClassGroup, id=class_id)
    
    # Dapatkan semua user yang memiliki role Wali Kelas atau Guru untuk ditampilkan di dropdown
    all_teachers = User.objects.filter(role__in=['WALIKELAS', 'GURU'])

    if request.method == 'POST':
        # --- INI BAGIAN LOGIKA PENYIMPANAN YANG PENTING ---

        # 1. Ambil data dari form
        class_name = request.POST.get('name')
        class_level = request.POST.get('level')
        # Ambil USERNAME wali kelas yang dipilih dari dropdown
        homeroom_teacher_username = request.POST.get('homeroom_teacher')

        # 2. Update field nama dan jenjang
        class_group.name = class_name
        class_group.level = class_level

        # 3. Cari objek User berdasarkan username yang dipilih
        try:
            # Jika ada username yang dipilih dari form
            if homeroom_teacher_username:
                teacher_object = User.objects.get(username=homeroom_teacher_username)
                # 4. Simpan OBJEK USER tersebut ke field homeroom_teacher
                class_group.homeroom_teacher = teacher_object
            else:
                # Jika tidak ada yang dipilih, set menjadi None (tidak ada wali kelas)
                class_group.homeroom_teacher = None
            
            # 5. Simpan semua perubahan ke database
            class_group.save()
            messages.success(request, f"Kelas '{class_group.name}' berhasil diperbarui.")
            return redirect('url_ke_daftar_kelas') # Ganti dengan URL Anda

        except User.DoesNotExist:
            # Tangani jika username yang dikirim tidak valid
            messages.error(request, f"Guru dengan username '{homeroom_teacher_username}' tidak ditemukan.")
        except Exception as e:
            messages.error(request, f"Terjadi kesalahan: {e}")

    # Jika metodenya GET (hanya menampilkan form)
    context = {
        'class_group': class_group,
        'all_teachers': all_teachers,
    }
    return render(request, 'dashboard/class_edit.html', context) # Ganti dengan path template Anda

@login_required
def grading_settings_view(request):
    """
    Menangani tampilan dan logika untuk mengubah pengaturan bobot penilaian.
    Hanya Admin yang bisa mengubah.
    """
    # Hanya Admin yang diizinkan mengubah pengaturan ini
    if request.user.role != 'ADMIN':
        messages.error(request, "Anda tidak memiliki izin untuk mengakses halaman ini.")
        return redirect('dashboard') # Arahkan ke halaman dashboard utama

    # Ambil pengaturan pertama (seharusnya hanya ada satu), atau buat jika belum ada.
    settings, created = GradingSettings.objects.get_or_create(pk=1)

    if request.method == 'POST':
        form = GradingSettingsForm(request.POST, instance=settings)
        if form.is_valid():
            form.save()
            messages.success(request, 'Pengaturan bobot penilaian berhasil diperbarui.')
            return redirect('grading_settings') # Kembali ke halaman yang sama
    else:
        form = GradingSettingsForm(instance=settings)

    context = {
        'form': form
    }
    return render(request, 'dashboard/grading_settings.html', context)

@login_required
def competency_settings_view(request, MataPelajaran_pk):
    MataPelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    # Pastikan guru yang login mengajar mata pelajaran ini
    # (memerlukan logika tambahan berdasarkan model penugasan guru Anda)

    if request.method == 'POST':
        # Logika untuk menyimpan kompetensi baru
        form = CompetencyForm(request.POST)
        if form.is_valid():
            new_competency = form.save(commit=False)
            new_competency.MataPelajaran = MataPelajaran
            new_competency.save()
            messages.success(request, 'Kompetensi baru berhasil ditambahkan.')
            return redirect('competency_settings', MataPelajaran_pk=MataPelajaran.pk)
    else:
        form = CompetencyForm()

    competencies = Competency.objects.filter(MataPelajaran=MataPelajaran)
    context = {
        'MataPelajaran': MataPelajaran,
        'competencies': competencies,
        'form': form,
    }
    return render(request, 'dashboard/competency_settings.html', context)

@login_required
def grade_input_view(request, class_pk, MataPelajaran_pk):
    class_group = get_object_or_404(ClassGroup, pk=class_pk)
    MataPelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    
    students = User.objects.filter(class_group=class_group, role='SISWA')
    competencies = Competency.objects.filter(MataPelajaran=MataPelajaran)

    if request.method == 'POST':
        # Logika kompleks untuk menyimpan puluhan nilai sekaligus
        for student in students:
            for competency in competencies:
                # Ambil nilai formatif dari form
                formative_value = request.POST.get(f'score_{student.pk}_{competency.pk}_FORMATIF')
                # Ambil nilai sumatif dari form
                summative_value = request.POST.get(f'score_{student.pk}_{competency.pk}_SUMATIF')

                # Simpan atau perbarui nilai formatif
                Score.objects.update_or_create(
                    student=student, competency=competency, score_type='FORMATIF',
                    defaults={'value': formative_value, 'teacher': request.user}
                )
                # Simpan atau perbarui nilai sumatif
                Score.objects.update_or_create(
                    student=student, competency=competency, score_type='SUMATIF',
                    defaults={'value': summative_value, 'teacher': request.user}
                )
        messages.success(request, 'Nilai berhasil disimpan.')
        return redirect('grade_input', class_pk=class_pk, MataPelajaran_pk=MataPelajaran_pk)

    # Siapkan data nilai yang sudah ada untuk ditampilkan di form
    scores_data = {}
    for student in students:
        scores_data[student.pk] = {
            (score.competency_id, score.score_type): score.value 
            for score in Score.objects.filter(student=student, competency__in=competencies)
        }

    context = {
        'class_group': class_group,
        'MataPelajaran': MataPelajaran,
        'students': students,
        'competencies': competencies,
        'scores_data': scores_data,
    }
    return render(request, 'dashboard/grade_input.html', context)
@login_required
def competency_settings_view(request, MataPelajaran_pk):
    MataPelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    
    # Di sini Anda bisa menambahkan validasi untuk memastikan hanya guru
    # yang mengajar mapel ini yang bisa mengakses.

    if request.method == 'POST':
        form = CompetencyForm(request.POST)
        if form.is_valid():
            # Simpan kompetensi baru, tapi kaitkan dulu dengan mata pelajaran yang benar
            new_competency = form.save(commit=False)
            new_competency.MataPelajaran = MataPelajaran
            try:
                new_competency.save()
                messages.success(request, f"Kompetensi '{new_competency.code}' berhasil ditambahkan.")
            except Exception as e:
                messages.error(request, f"Gagal menyimpan. Kode kompetensi mungkin sudah ada. Error: {e}")
            
            return redirect('competency_settings', MataPelajaran_pk=MataPelajaran.pk)
    else:
        form = CompetencyForm()

    # Ambil semua kompetensi yang sudah ada untuk mata pelajaran ini
    competencies = Competency.objects.filter(MataPelajaran=MataPelajaran).order_by('code')
    
    context = {
        'MataPelajaran': MataPelajaran,
        'competencies': competencies,
        'form': form,
    }
    return render(request, 'dashboard/competency_settings.html', context)

@login_required
def grade_input_view(request, class_pk, MataPelajaran_pk):
    """
    Menampilkan halaman untuk menginput nilai formatif dan sumatif
    untuk semua siswa di sebuah kelas pada mata pelajaran tertentu.
    """
    class_group = get_object_or_404(ClassGroup, pk=class_pk)
    MataPelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    
    # Ambil semua siswa di kelas tersebut
    students = User.objects.filter(class_group=class_group, role='SISWA').order_by('first_name')
    # Ambil semua kompetensi untuk mata pelajaran tersebut
    competencies = Competency.objects.filter(MataPelajaran=MataPelajaran).order_by('code')

    if request.method == 'POST':
        # Logika untuk menyimpan semua nilai yang diinput dari form
        for student in students:
            for competency in competencies:
                # Ambil nilai formatif dari form POST
                formative_value = request.POST.get(f'score_{student.pk}_{competency.pk}_FORMATIF')
                # Ambil nilai sumatif dari form POST
                summative_value = request.POST.get(f'score_{student.pk}_{competency.pk}_SUMATIF')

                # Simpan atau perbarui nilai formatif jika ada isinya
                if formative_value:
                    Score.objects.update_or_create(
                        student=student, competency=competency, score_type='FORMATIF',
                        defaults={'value': formative_value, 'teacher': request.user}
                    )
                
                # Simpan atau perbarui nilai sumatif jika ada isinya
                if summative_value:
                    Score.objects.update_or_create(
                        student=student, competency=competency, score_type='SUMATIF',
                        defaults={'value': summative_value, 'teacher': request.user}
                    )

        messages.success(request, 'Semua perubahan nilai berhasil disimpan.')
        return redirect('grade_input', class_pk=class_pk, MataPelajaran_pk=MataPelajaran_pk)

    # Siapkan data nilai yang sudah ada untuk ditampilkan di form
    # Ini adalah dictionary untuk akses cepat di template
    scores_data = {}
    all_scores = Score.objects.filter(student__in=students, competency__in=competencies)
    for score in all_scores:
        if score.student_id not in scores_data:
            scores_data[score.student_id] = {}
        if score.competency_id not in scores_data[score.student_id]:
            scores_data[score.student_id][score.competency_id] = {}
        scores_data[score.student_id][score.competency_id][score.score_type] = score.value

    context = {
        'class_group': class_group,
        'MataPelajaran': MataPelajaran,
        'students': students,
        'competencies': competencies,
        'scores_data': scores_data,
    }
    return render(request, 'dashboard/grade_input.html', context)
@login_required
def competency_settings_view(request, MataPelajaran_pk):
    MataPelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    
    # Membuat FormSet untuk AssessmentComponent yang terikat pada Competency
    # Ini memungkinkan kita mengedit beberapa komponen sekaligus
    ComponentFormSet = inlineformset_factory(
        Competency, 
        AssessmentComponent, 
        fields=('name', 'description', 'weight'),
        extra=1, # Tampilkan 1 form kosong secara default
        can_delete=True # Izinkan penghapusan komponen
    )

    if request.method == 'POST':
        # Untuk saat ini, kita fokus pada pembuatan kompetensi baru
        # (Logika edit bisa ditambahkan nanti)
        form = CompetencyForm(request.POST)
        if form.is_valid():
            new_competency = form.save(commit=False)
            new_competency.MataPelajaran = MataPelajaran
            
            formset = ComponentFormSet(request.POST, instance=new_competency)
            if formset.is_valid():
                new_competency.save()
                formset.save()
                messages.success(request, f"Kompetensi '{new_competency.code}' dan komponennya berhasil disimpan.")
                return redirect('competency_settings', MataPelajaran_pk=MataPelajaran.pk)
            else:
                messages.error(request, f"Terdapat error pada data komponen: {formset.errors}")
        else:
            messages.error(request, f"Terdapat error pada data kompetensi: {form.errors}")

    # Logika untuk menampilkan halaman
    form = CompetencyForm()
    formset = ComponentFormSet()
    
    all_competencies = Competency.objects.filter(MataPelajaran=MataPelajaran).prefetch_related('components')

    context = {
        'MataPelajaran': MataPelajaran,
        'all_competencies': all_competencies,
        'form': form,
        'formset': formset,
    }
    return render(request, 'dashboard/competency_settings_advanced.html', context)

@login_required
def question_create_view(request, exam_pk):
    exam = get_object_or_404(Exam, pk=exam_pk)
    
    # Membuat FormSet dari QuestionForm, kita minta 1 form awal
    QuestionFormSet = formset_factory(QuestionForm, extra=1)

    if request.method == 'POST':
        formset = QuestionFormSet(request.POST)
        if formset.is_valid():
            created_count = 0
            for form in formset:
                if form.has_changed(): # Hanya proses form yang diisi
                    question = form.save(commit=False)
                    question.exam = exam
                    question.MataPelajaran = exam.MataPelajaran
                    question.save()
                    created_count += 1
            
            if created_count > 0:
                messages.success(request, f'{created_count} soal baru berhasil ditambahkan!')
            return redirect('exam_detail', pk=exam.pk)
    else:
        form = QuestionForm()
    
    context = {
        'form': form,
        'exam': exam,
        'page_title': 'Buat Soal Baru'
    }
    return render(request, 'dashboard/question_form.html', context)

# View untuk edit (kemungkinan Anda sudah punya)
login_required
@login_required
def question_update_view(request, pk):
    question = get_object_or_404(Question, pk=pk)
    
    # Menentukan exam dari question untuk konteks, jika ada
    # Ini penting untuk ditampilkan di template
    exam = question.exam if hasattr(question, 'exam') else None

    if request.method == 'POST':
        form = QuestionForm(request.POST, request.FILES, instance=question)
        if form.is_valid():
            # Simpan form terlebih dahulu
            form.save()
            messages.success(request, 'Soal berhasil diperbarui.')
            
            # =================================================================
            # PERBAIKAN UTAMA ADA DI SINI
            # =================================================================
            # Setelah disimpan, kita perlu memeriksa kembali instance 'question'
            # yang sudah diperbarui.

            # Cek apakah soal ini memiliki relasi ke sebuah ujian.
            if question.exam:
                # Jika YA, redirect ke halaman detail ujian tersebut.
                return redirect('exam_detail', pk=question.exam.pk)
            else:
                # Jika TIDAK (misal, soal dari Bank Soal), 
                # redirect ke halaman manajemen soal utama.
                return redirect('question_management') # Pastikan Anda punya URL name ini
    else:
        form = QuestionForm(instance=question)
    
    context = {
        'form': form,
        'page_title': f"Edit Soal #{question.pk}",
        'exam': exam  # Kirim variabel exam ke template
    }
    return render(request, 'dashboard/question_form.html', context)


@login_required
def redirect_view(request):
    role = request.user.role
    if role == 'ADMIN':
        return redirect('dashboard_admin')
    elif role in ['GURU', 'WALIKELAS']:
        return redirect('dashboard_guru')
    elif role == 'SISWA':
        return redirect('dashboard_siswa')
    elif role == 'MANAGER':
        return redirect('dashboard_manager')
    else:
        return redirect('profile')

@login_required
def rapor_print_view(request):
    """
    View ini hanya bertugas untuk menampilkan halaman utama dashboard rapor (rapor_print.html).
    Semua data akan dimuat secara dinamis oleh JavaScript.
    """
    return render(request, "dashboard/rapor_print.html")


@login_required
def raport_data_api(request):
    """
    API untuk menyediakan data rapor siswa lengkap berdasarkan wali kelas yang login.
    """
    wali_kelas = request.user

    try:
        daftar_siswa = User.objects.filter(
            role='SISWA',
            class_group__homeroom_teacher=wali_kelas
        )

        info_sekolah = InfoSekolah.objects.first()
        if not info_sekolah:
            info_sekolah = InfoSekolah(nama="Nama Sekolah Belum Diatur", alamat="-", kepala_sekolah="-", nip_kepsek="-")

        students_json_data = []
        for siswa in daftar_siswa:
            # --- Nilai Mata Pelajaran ---
            nilai_list = []
            for idx, n in enumerate(siswa.nilai_siswa.all().order_by('mata_pelajaran__nama')):
                if n.mata_pelajaran:
                    try:
                        nilai_angka = float(n.nilai_akhir)
                        nilai_list.append({
                            'no': idx + 1,
                            'nama': n.mata_pelajaran.nama,
                            'nilai': nilai_angka,
                            'deskripsi': n.deskripsi
                        })
                    except (ValueError, TypeError):
                        continue

            # --- Ekstrakurikuler ---
            ekskul_list = [
                {
                    'nama': e.nama_kegiatan,
                    'predikat': e.predikat or "-",
                    'keterangan': e.keterangan or "-"
                }
                for e in siswa.ekskul_siswa.all()
            ]

            # --- Prestasi ---
            prestasi_list = [
                {
                    'jenis': p.jenis_kegiatan,
                    'keterangan': p.keterangan
                }
                for p in siswa.prestasi.all()
            ]

            # --- Kehadiran ---
            try:
                kehadiran_obj = siswa.kehadiran_siswa
                kehadiran_data = {
                    'sakit': f"{kehadiran_obj.sakit} Hari",
                    'izin': f"{kehadiran_obj.izin} Hari",
                    'tanpaKeterangan': f"{kehadiran_obj.alpha} Hari",
                }
            except ObjectDoesNotExist:
                kehadiran_data = {
                    'sakit': "0 Hari", 'izin': "0 Hari", 'tanpaKeterangan': "0 Hari"
                }

            # --- Catatan Wali Kelas ---
            catatan_obj = siswa.catatan_wali_kelas.first()
            catatan_text = catatan_obj.catatan if catatan_obj else ""

            # --- Tanggapan Orang Tua ---
            tanggapan_obj = siswa.tanggapan_orang_tua.last()
            tanggapan_text = tanggapan_obj.tanggapan if tanggapan_obj else ""

            # --- Kenaikan Kelas ---
            kenaikan_obj = siswa.kenaikan_kelas.last()
            kenaikan_status = {
                'status': kenaikan_obj.get_status_display(),
                'keterangan': kenaikan_obj.keterangan
            } if kenaikan_obj else {}

            # --- Susun Data ---
            student_data = {
                'id': str(siswa.id),
                'nama': siswa.get_full_name() or siswa.username,
                'kelas': siswa.class_group.name if siswa.class_group else "-",
                'nisn': f"{siswa.nis or '-'}/{siswa.nisn or '-'}",
                'semester': 'II/Genap',
                'fase': 'E',
                'tahunPelajaran': '2024/2025',
                'infoSekolah': {
                    'nama': info_sekolah.nama, 'alamat': info_sekolah.alamat,
                    'kepalaSekolah': info_sekolah.kepala_sekolah, 'nipKepsek': info_sekolah.nip_kepsek,
                },
                'infoWaliKelas': {
                    'nama': wali_kelas.get_full_name() or wali_kelas.username,
                    'nip': wali_kelas.employee_id or '-',
                },
                'matapelajaran': nilai_list,
                'ekstrakurikuler': ekskul_list,
                'prestasi': prestasi_list,
                'ketidakhadiran': kehadiran_data,
                'catatanWaliKelas': catatan_text,
                'tanggapanOrangTua': tanggapan_text,
                'kenaikanKelas': kenaikan_status,
            }
            students_json_data.append(student_data)

        return JsonResponse(students_json_data, safe=False)

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_POST
def update_rapor_api(request):
    """
    API untuk menyimpan perubahan data rapor (nilai, ekskul, prestasi, catatan, tanggapan, kenaikan kelas).
    """
    try:
        data = json.loads(request.body)
        student_id = data.get('studentId')

        siswa = User.objects.get(id=student_id, role='SISWA', class_group__homeroom_teacher=request.user)

        # --- Update Nilai Mata Pelajaran ---
        for mapel_data in data.get('mataPelajaran', []):
            mapel_obj = MataPelajaran.objects.get(nama=mapel_data['nama'])
            Nilai.objects.update_or_create(
                siswa=siswa,
                mata_pelajaran=mapel_obj,
                defaults={
                    'nilai_akhir': mapel_data['nilai'],
                    'deskripsi': mapel_data['deskripsi']
                }
            )

        # --- Update Ekstrakurikuler ---
        for ekskul_data in data.get('ekstrakurikuler', []):
            Ekstrakurikuler.objects.update_or_create(
                siswa=siswa,
                nama_kegiatan=ekskul_data['nama'],
                defaults={
                    'predikat': ekskul_data.get('predikat'),
                    'keterangan': ekskul_data.get('keterangan')
                }
            )

        # --- Update Prestasi ---
        for prestasi_data in data.get('prestasi', []):
            Prestasi.objects.update_or_create(
                siswa=siswa,
                jenis_kegiatan=prestasi_data['jenis'],
                defaults={'keterangan': prestasi_data.get('keterangan')}
            )

        # --- Update Catatan Wali Kelas ---
        CatatanWaliKelas.objects.update_or_create(
            siswa=siswa,
            defaults={'catatan': data.get('catatanWaliKelas', '')}
        )

        # --- Tambah Tanggapan Orang Tua ---
        if 'tanggapanOrangTua' in data:
            TanggapanOrangTua.objects.create(
                siswa=siswa,
                tanggapan=data['tanggapanOrangTua']
            )

        # --- Update Kenaikan Kelas ---
        if 'kenaikanKelas' in data:
            kenaikan_data = data['kenaikanKelas']
            KenaikanKelas.objects.update_or_create(
                siswa=siswa,
                defaults={
                    'status': kenaikan_data.get('status'),
                    'keterangan': kenaikan_data.get('keterangan', '')
                }
            )

        return JsonResponse({'status': 'success', 'message': 'Data rapor berhasil disimpan!'})

    except ObjectDoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Siswa tidak ditemukan atau Anda tidak berhak mengedit.'}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
def lampiran_print_view(request):
    """
    View ini khusus untuk halaman "Cetak Lampiran".
    Ia akan mencari siswa berdasarkan wali kelas yang login dan mengirimkan
    data tersebut langsung ke template lampiran_print.html.
    """
    wali_kelas = request.user
    
    # Lakukan query untuk mencari semua siswa yang kelasnya diampu oleh wali kelas ini
    daftar_siswa = User.objects.filter(
        role='SISWA', 
        class_group__homeroom_teacher=wali_kelas
    )
    
    context = {
        'siswa': daftar_siswa
    }
    
    return render(request, 'dashboard/lampiran_print.html', context)

@login_required
def dashboard_admin(request):
    return render(request, "dashboard/dashboard_admin.html")

@login_required
def dashboard_guru(request):
    return render(request, "dashboard/dashboard_guru.html")

@login_required
def dashboard_siswa(request):
    return render(request, "dashboard/dashboard_siswa.html")

@login_required
def dashboard_manager(request):
    return render(request, "dashboard/dashboard_manager.html")

@login_required
def import_questions_for_MataPelajaran_view(request, MataPelajaran_pk):
    """
    Menangani impor soal dari file Word (.docx) untuk MataPelajaran TERTENTU.
    Fungsi ini sekarang mendukung EKSTRAKSI GAMBAR dari sel pertanyaan.
    """
    mata_pelajaran = get_object_or_404(MataPelajaran, pk=MataPelajaran_pk)
    
    if request.method == 'POST' and request.FILES.get('word_file'):
        word_file = request.FILES['word_file']
        if not word_file.name.endswith('.docx'):
            messages.error(request, 'Format file tidak valid. Harap unggah file Word (.docx).')
            return redirect('question_management')

        try:
            document = docx.Document(word_file)
            if not document.tables:
                messages.error(request, "File Word yang diunggah tidak berisi tabel.")
                return redirect('question_management')

            table = document.tables[0]
            created_count = 0
            
            # --- KAMUS PEMETAAN UNTUK DIFFICULTY ---
            # PENTING: Sesuaikan kunci (kiri) dengan apa yang tertulis di file Word Anda (case-insensitive)
            # dan value (kanan) dengan nilai di model DifficultyLevel Anda.
            difficulty_map = {
                'prastruktural': 'PRA-STRUKTURAL',
                'unistruktural': 'UNISTRUKTURAL',
                'multistruktural': 'MULTISTRUKTURAL',
                'relasi': 'RELASI',
                'abstrak': 'ABSTRAK',
                # Tambahkan pemetaan lain jika perlu, contoh:
                # 'sangat sulit': 'RELASI', 
            }
            # -----------------------------------------

            for i, row in enumerate(table.rows[1:]):  # skip header
                try:
                    if not any(cell.text.strip() for cell in row.cells):
                        continue

                    # --- LOGIKA BARU UNTUK EKSTRAK GAMBAR DAN TEKS ---
                    pertanyaan_cell = row.cells[3]
                    pertanyaan_text = ""
                    image_blob = None

                    # Iterasi melalui setiap paragraf di dalam sel untuk mengambil teks dan gambar
                    for paragraph in pertanyaan_cell.paragraphs:
                        pertanyaan_text += paragraph.text + "\n"
                        # Cari gambar di dalam 'run' di paragraf
                        for run in paragraph.runs:
                            # Cek apakah ada elemen gambar (drawing)
                            if run.element.xpath('.//w:drawing'):
                                for drawing in run.element.xpath('.//w:drawing'):
                                    # Dapatkan ID hubungan gambar
                                    rId = drawing.find('.//a:blip', namespaces=drawing.nsmap).attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed']
                                    # Dapatkan bagian gambar dari dokumen
                                    image_part = document.part.related_parts[rId]
                                    image_blob = image_part.blob # Ini adalah data biner gambar
                                    break # Ambil gambar pertama yang ditemukan
                            if image_blob: break
                        if image_blob: break
                    
                    pertanyaan_text = pertanyaan_text.strip()
                    # ---------------------------------------------------

                    # Ambil nilai lain dari tabel
                    level_from_doc = row.cells[0].text.strip().lower() # ubah ke lowercase agar cocok dengan map
                    
                    # --- LOGIKA BARU UNTUK MENGGUNAKAN PEMETAAN ---
                    # Gunakan nilai default jika tidak ditemukan di map
                    level = difficulty_map.get(level_from_doc, 'PRA-STRUKTURAL') 
                    # ----------------------------------------------

                    bobot_text = row.cells[1].text.strip()
                    bobot = int(bobot_text) if bobot_text.isdigit() else 10
                    
                    tipe_soal = row.cells[2].text.strip().upper()
                    jawaban = row.cells[4].text.strip()

                    options_dict = {}
                    correct_answer = jawaban

                    # Logika parsing berdasarkan tipe soal
                    if tipe_soal == "PILIHAN_GANDA":
                        options_dict = {
                            label: row.cells[5 + j].text.strip()
                            for j, label in enumerate(['A', 'B', 'C', 'D', 'E'])
                            if (5 + j) < len(row.cells) and row.cells[5 + j].text.strip()
                        }

                    elif tipe_soal == "MENJODOHKAN":
                        left_items = [item.strip() for item in row.cells[5].text.split(";") if item.strip()]
                        right_items = [item.strip() for item in row.cells[6].text.split(";") if item.strip()]
                        options_dict = {"left": left_items, "right": right_items}
                        correct_answer = {
                            pair.split("-")[0].strip(): pair.split("-")[1].strip()
                            for pair in jawaban.split(";") if "-" in pair
                        }

                    elif tipe_soal == "ESAI":
                        options_dict = {}
                        correct_answer = jawaban

                    # Buat objek Question tanpa gambar terlebih dahulu
                    new_question = Question(
                        MataPelajaran=mata_pelajaran,
                        question_text=pertanyaan_text,
                        question_type=tipe_soal,
                        difficulty=level,
                        weight=bobot,
                        options=options_dict,
                        correct_answer=correct_answer,
                        exam=None
                    )

                    # --- LOGIKA BARU UNTUK MENYIMPAN GAMBAR ---
                    if image_blob:
                        # Buat nama file yang unik untuk menghindari tumpukan nama
                        file_name = f"soal_{uuid.uuid4().hex[:8]}.png"
                        # Simpan gambar ke field 'image'
                        new_question.image.save(file_name, ContentFile(image_blob), save=False)
                    # -------------------------------------------
                    
                    # Simpan objek Question ke database
                    new_question.save()
                    created_count += 1

                except Exception as e:
                    messages.warning(request, f"Gagal memproses baris {i + 2}: {e}")
            
            if created_count > 0:
                messages.success(request, f'{created_count} soal untuk mapel {mata_pelajaran.name} berhasil diimpor.')

        except Exception as e:
            messages.error(request, f'Terjadi error saat membaca file: {e}')

    return redirect('question_management')


@login_required
def export_MataPelajarans_template_view(request):
    """
    View ini menangani permintaan untuk men-download template Excel (.xlsx)
    kosong untuk data mata pelajaran.
    """
    # 1. Buat HttpResponse dengan header Excel yang benar.
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="template_mata_pelajaran.xlsx"'},
    )

    # 2. Buat Workbook dan worksheet Excel.
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = 'Data Mata Pelajaran'

    # 3. Tulis baris header.
    headers = ['Name', 'Level', 'Description']
    worksheet.append(headers)

    # 4. Simpan workbook ke dalam response.
    workbook.save(response)

    # 5. Kembalikan response sebagai file yang akan di-download.
    return response

@login_required
def item_analysis_view(request, exam_pk):
    """
    Menghitung dan menampilkan analisis butir soal untuk sebuah ujian.
    """
    exam = get_object_or_404(Exam, pk=exam_pk)
    questions = exam.questions.all().order_by('id')
    
    # Hitung total siswa yang sudah menyelesaikan ujian ini
    total_attempts = ExamAttempt.objects.filter(exam=exam, is_completed=True).count()
    
    analysis_results = []
    for q in questions:
        # Ambil semua jawaban untuk soal ini dari semua percobaan
        answers = Answer.objects.filter(question=q, attempt__exam=exam, attempt__is_completed=True)
        
        correct = answers.filter(is_correct=True).count()
        incorrect = answers.filter(is_correct=False).count()
        
        # Analisis opsi jawaban (hanya untuk Pilihan Ganda)
        option_analysis = {}
        if q.question_type == 'PILIHAN_GANDA' and q.options:
            try:
                # =======================================================
                # PERBAIKAN: Ubah string JSON menjadi dictionary Python
                # =======================================================
                options_dict = json.loads(q.options)
                
                # Menghitung berapa kali setiap opsi dipilih
                option_counts = answers.values('selected_option').annotate(count=Count('selected_option'))
                
                # Loop melalui dictionary yang sudah di-parsing
                for option_key in options_dict.keys():
                    count_data = next((item for item in option_counts if item["selected_option"] == option_key), None)
                    option_analysis[option_key] = count_data['count'] if count_data else 0

            except (json.JSONDecodeError, TypeError):
                # Jika format 'options' salah, lewati analisis opsi
                option_analysis = {"error": "Format opsi jawaban tidak valid."}

        analysis_results.append({
            'question': q,
            'correct': correct,
            'incorrect': incorrect,
            'unanswered': total_attempts - (correct + incorrect),
            'option_analysis': option_analysis
        })

    context = {
        'exam': exam,
        'results': analysis_results,
        'total_attempts': total_attempts,
        'page_title': f'Analisis Butir Soal: {exam.name}'
    }
    
    return render(request, 'dashboard/item_analysis.html', context)

@login_required
def export_item_analysis_excel(request, exam_pk):
    """
    Menghasilkan file Excel (.xlsx) dari data analisis butir soal.
    """
    exam = get_object_or_404(Exam, pk=exam_pk)
    questions = exam.questions.all().order_by('id')
    total_attempts = ExamAttempt.objects.filter(exam=exam, is_completed=True).count()

    # Logika untuk mengambil data analisis (sama seperti di item_analysis_view)
    analysis_results = []
    for q in questions:
        answers = Answer.objects.filter(question=q, attempt__exam=exam)
        correct = answers.filter(is_correct=True).count()
        incorrect = answers.filter(is_correct=False).count()
        option_analysis = {}
        if q.question_type == 'PILIHAN_GANDA' and q.options:
            option_counts = answers.values('selected_option').annotate(count=Count('selected_option'))
            for option_key in q.options.keys():
                count_data = next((item for item in option_counts if item["selected_option"] == option_key), None)
                option_analysis[option_key] = count_data['count'] if count_data else 0
        analysis_results.append({ 'question': q, 'correct': correct, 'incorrect': incorrect, 'unanswered': total_attempts - (correct + incorrect), 'option_analysis': option_analysis })

    # Buat response Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Analisis Soal - {exam.name}.xlsx"'

    # Buat Workbook dan Worksheet
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Analisis Butir Soal"

    # Style
    bold_font = Font(bold=True)
    center_align = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    # Menulis Header Utama
    ws.merge_cells('A1:K1')
    ws['A1'] = f"Analisis Butir Soal: {exam.name}"
    ws['A1'].font = bold_font
    ws['A1'].alignment = center_align

    # Menulis Header Tabel
    headers = ["No.", "Teks Soal", "Tipe", "Benar", "Salah", "Kosong", "Opsi A", "Opsi B", "Opsi C", "Opsi D", "Opsi E"]
    for col_num, header_title in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num, value=header_title)
        cell.font = bold_font
        cell.border = thin_border

    # Menulis Data Soal
    current_row = 4
    for item in analysis_results:
        ws.cell(row=current_row, column=1, value=current_row - 3).border = thin_border
        ws.cell(row=current_row, column=2, value=item['question'].question_text).border = thin_border
        ws.cell(row=current_row, column=3, value=item['question'].get_question_type_display()).border = thin_border
        ws.cell(row=current_row, column=4, value=item['correct']).border = thin_border
        ws.cell(row=current_row, column=5, value=item['incorrect']).border = thin_border
        ws.cell(row=current_row, column=6, value=item['unanswered']).border = thin_border
        
        # Opsi Pilihan Ganda
        options = item.get('option_analysis', {})
        ws.cell(row=current_row, column=7, value=options.get('A', '')).border = thin_border
        ws.cell(row=current_row, column=8, value=options.get('B', '')).border = thin_border
        ws.cell(row=current_row, column=9, value=options.get('C', '')).border = thin_border
        ws.cell(row=current_row, column=10, value=options.get('D', '')).border = thin_border
        ws.cell(row=current_row, column=11, value=options.get('E', '')).border = thin_border
        
        current_row += 1

    # Atur lebar kolom
    ws.column_dimensions['B'].width = 50
    
    # Simpan workbook ke response
    wb.save(response)
    return response
@login_required
def correction_detail_view(request, attempt_pk):
    attempt = get_object_or_404(ExamAttempt.objects.select_related('student', 'exam'), pk=attempt_pk)
    answers = Answer.objects.filter(attempt=attempt).select_related('question').order_by('question__id')

    # Hitung skor awal hanya dari soal Pilihan Ganda yang sudah dinilai otomatis
    pg_answers = answers.filter(question__question_type='PILIHAN_GANDA')
    correct_pg_count = pg_answers.filter(is_correct=True).count()
    total_pg_questions = attempt.exam.questions.filter(question_type='PILIHAN_GANDA').count()
    
    initial_score = 0
    if total_pg_questions > 0:
        initial_score = (correct_pg_count / total_pg_questions) * 100

    context = {
        'attempt': attempt,
        'answers': answers,
        'initial_score': initial_score,
    }
    return render(request, 'dashboard/teacher/correction_detail.html', context)
@login_required
def get_ai_correction_suggestion_view(request, answer_pk):
    answer = get_object_or_404(Answer, pk=answer_pk)
    
    # Konfigurasi API
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-pro')
    except Exception as e:
        return JsonResponse({'error': f'Konfigurasi AI gagal: {e}'}, status=500)

    # Prompt Engineering: Membentuk pertanyaan yang baik untuk AI
    prompt = f"""
    Anda adalah seorang asisten guru ahli yang objektif. 
    Tugas Anda adalah memberikan saran skor dan analisis singkat untuk jawaban esai siswa.
    
    Konteks:
    - Pertanyaan: "{answer.question.question_text}"
    - Kunci Jawaban (sebagai referensi): "{answer.question.correct_answer}"
    - Bobot / Skor Maksimal untuk soal ini: {answer.question.weight}
    
    Jawaban Siswa:
    "{answer.answer_text}"
    
    Tugas Anda:
    Berdasarkan konteks di atas, berikan evaluasi atas jawaban siswa. 
    Jawaban Anda HARUS dalam format JSON yang valid dengan dua kunci: "suggested_score" (angka antara 0 dan skor maksimal) dan "analysis" (string singkat berisi justifikasi skor Anda, sebutkan poin positif dan negatif dari jawaban siswa).
    Contoh JSON: {{"suggested_score": 8, "analysis": "Jawaban siswa sudah cukup baik dan mencakup poin utama, namun kurang detail pada bagian penjelasan."}}
    
    Hanya berikan JSON, tanpa teks tambahan sebelum atau sesudahnya.
    """

    try:
        response = model.generate_content(prompt)
        ai_response_text = response.text.strip()
        
        # Membersihkan dan parsing JSON
        ai_data = json.loads(ai_response_text)
        
        # Menambahkan skor maksimal ke response untuk ditampilkan di UI
        ai_data['max_score'] = answer.question.weight

        return JsonResponse(ai_data)
    except Exception as e:
        return JsonResponse({'error': f'Gagal mendapatkan respon dari AI: {e}'}, status=500)
@login_required
def user_edit_view(request, user_id): # Nama view mungkin berbeda di kode Anda
    """
    View untuk menangani pengeditan data pengguna.
    """
    user_instance = get_object_or_404(User, id=user_id)
    
    if request.method == 'POST':
        form = UserEditForm(request.POST, instance=user_instance)
        if form.is_valid():
            form.save()
            messages.success(request, f"Data pengguna '{user_instance.username}' berhasil diperbarui.")
            return redirect('user_management') # Ganti dengan URL halaman manajemen pengguna Anda
    else:
        form = UserEditForm(instance=user_instance)

    context = {
        'form': form,
        'user_instance': user_instance,
    }
    return render(request, 'dashboard/user_edit.html', context) # Ganti dengan path template Anda
@login_required
def MataPelajaran_management_view(request):
    """
    View untuk menampilkan dan menangani halaman Manajemen Mata Pelajaran.
    """
    # Logika untuk menangani semua permintaan POST
    if request.method == 'POST':
        # Cek apakah ini permintaan HAPUS MASSAL
        if 'delete_selected' in request.POST:
            MataPelajaran_ids_to_delete = request.POST.getlist('MataPelajaran_ids')
            if MataPelajaran_ids_to_delete:
                MataPelajarans_to_delete = MataPelajaran.objects.filter(pk__in=MataPelajaran_ids_to_delete)
                count = MataPelajarans_to_delete.count()
                MataPelajarans_to_delete.delete()
                messages.success(request, f'{count} mata pelajaran berhasil dihapus.')
            else:
                messages.warning(request, 'Tidak ada mata pelajaran yang dipilih untuk dihapus.')
            # Ganti 'MataPelajaran_management' dengan nama URL Anda jika berbeda
            return redirect('MataPelajaran_management')

        # Jika bukan, anggap ini permintaan TAMBAH MAPEL BARU
        else:
            form = MataPelajaranForm(request.POST)
            if form.is_valid():
                form.save()
                messages.success(request, 'Mata pelajaran baru berhasil ditambahkan.')
                return redirect('MataPelajaran_management')
            # Jika form tidak valid, biarkan proses lanjut ke bawah agar error ditampilkan
    
    # Logika untuk permintaan GET (menampilkan halaman) atau form yang tidak valid
    form = MataPelajaranForm()
    
    # --- PERBAIKAN DI SINI ---
    # 1. Ambil semua mata pelajaran
    MataPelajarans = MataPelajaran.objects.all().order_by('nama')
    
    # 2. Ambil daftar jenjang yang unik dari mata pelajaran yang ada
    levels = MataPelajaran.objects.values_list('level', flat=True).distinct().order_by('level')
    # Filter untuk menghapus None atau string kosong jika ada
    existing_levels = [level for level in levels if level]
    # --- AKHIR PERBAIKAN ---

    assignments = TeacherAssignment.objects.all().order_by('teacher__first_name', 'class_group__name')
    
    context = {
       'form': form,
        'MataPelajarans': MataPelajarans,
        'levels': existing_levels, # Kirim daftar jenjang ke template
        'assignments': assignments,
    }
    # Ganti dengan path template Anda jika berbeda
    return render(request, 'dashboard/MataPelajaran_management.html', context)
@login_required
def export_assignments_template_view(request):
    """
    Membuat dan mengirimkan file template Excel untuk penugasan guru.
    """
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="template_penugasan_guru.xlsx"'},
    )
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = 'Penugasan Guru'
    
    # Header ini harus cocok dengan file template Anda
    headers = ['teacher_username', 'MataPelajaran_name', 'MataPelajaran_level', 'class_name']
    worksheet.append(headers)
    
    workbook.save(response)
    return response
@login_required
def import_assignments_view(request):
    """
    View untuk menangani upload dan impor data penugasan guru dari file Excel.
    """
    # Ganti 'assignment_management' dengan nama URL halaman sebelumnya jika berbeda
    redirect_url = 'assignment_management' 

    if request.method == 'POST':
        excel_file = request.FILES.get('excel_file')
        if not excel_file:
            messages.error(request, "Tidak ada file yang diunggah.")
            return redirect(redirect_url)

        try:
            workbook = openpyxl.load_workbook(excel_file)
            worksheet = workbook.active
            
            # Verifikasi header
            expected_headers = ['teacher_username', 'MataPelajaran_name', 'MataPelajaran_level', 'class_name']
            actual_headers = [cell.value for cell in worksheet[1]]
            if expected_headers != actual_headers:
                messages.error(request, f"Format file tidak sesuai. Harap gunakan template dengan header: {', '.join(expected_headers)}")
                return redirect(redirect_url)

            errors = []
            success_count = 0
            for i, row in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
                teacher_username, MataPelajaran_name, MataPelajaran_level, class_name = row
                
                try:
                    teacher = User.objects.get(username=teacher_username)
                    MataPelajaran = MataPelajaran.objects.get(nama=MataPelajaran_name)
                    class_group = ClassGroup.objects.get(name=class_name)
                    
                    _, created = TeacherAssignment.objects.update_or_create(
                        teacher=teacher,
                        MataPelajaran=MataPelajaran,
                        class_group=class_group
                    )
                    success_count += 1

                except User.DoesNotExist:
                    errors.append(f"Baris {i}: Guru '{teacher_username}' tidak ditemukan.")
                except MataPelajaran.DoesNotExist:
                    errors.append(f"Baris {i}: Mata Pelajaran '{MataPelajaran_name}' tidak ditemukan.")
                except ClassGroup.DoesNotExist:
                    errors.append(f"Baris {i}: Kelas '{class_name}' tidak ditemukan.")
                except Exception as e:
                    errors.append(f"Baris {i}: Terjadi error - {e}")

            if success_count > 0:
                messages.success(request, f"{success_count} data penugasan berhasil diimpor/diperbarui.")
            if errors:
                for error in errors:
                    messages.warning(request, error)

        except Exception as e:
            messages.error(request, f"Gagal memproses file: {e}")
        
        return redirect(redirect_url)
    
    return redirect(redirect_url)
@login_required
def theme_settings(request):
    user = request.user
    if request.method == "POST":
        form = ThemeForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            return redirect('role_redirect')  # kembali ke dashboard setelah simpan
    else:
        form = ThemeForm(instance=user)

    return render(request, 'dashboard/theme_settings.html', {'form': form})
    
def admin_dashboard(request):
    total_users = User.objects.count()
    total_guru = User.objects.filter(role=User.Role.GURU).count()
    total_siswa = User.objects.filter(role=User.Role.SISWA).count()
    total_walikelas = User.objects.filter(role=User.Role.WALIKELAS).count()
    total_manager = User.objects.filter(role=User.Role.MANAGER).count()
    total_admin = User.objects.filter(role=User.Role.ADMIN).count()

    # Statistik bulanan (user join per bulan)
    user_per_month = (
        User.objects.annotate(month=TruncMonth("date_joined"))
        .values("month")
        .annotate(count=Count("id"))
        .order_by("month")
    )

    months = [u["month"].strftime("%b %Y") for u in user_per_month]
    month_counts = [u["count"] for u in user_per_month]

    context = {
        "total_users": total_users,
        "total_guru": total_guru,
        "total_siswa": total_siswa,
        "total_walikelas": total_walikelas,
        "total_manager": total_manager,
        "total_admin": total_admin,
        "months": months,
        "month_counts": month_counts,
    }
    return render(request, "dashboard/dashboard_admin.html", context)
@login_required
def manajemen_nilai_view(request):
    """
    View ini bertanggung jawab untuk merender halaman awal manajemen nilai.
    """
    daftar_kelas = []
    daftar_mapel = []

    try:
        daftar_kelas = ClassGroup.objects.all().order_by('name')
    except Exception as e:
        logger.error(f"Error saat mengambil ClassGroup: {e}")

    try:
        daftar_mapel = MataPelajaran.objects.all().order_by('nama')
    except Exception as e:
        logger.error(f"Error saat mengambil MataPelajaran: {e}")

    context = {
        'daftar_kelas': daftar_kelas,
        'daftar_mapel': daftar_mapel,
        'title': 'Manajemen Nilai dan Rapor'
    }
    
    return render(request, 'dashboard/manajemen_nilai.html', context)
@user_passes_test(is_admin) # Hanya admin yang bisa akses halaman ini
def komponen_nilai_list(request):
    komponen_list = KomponenNilai.objects.all()
    form = KomponenNilaiForm()

    if request.method == 'POST':
        form = KomponenNilaiForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('komponen_nilai_list') # Kembali ke halaman yang sama setelah berhasil

    context = {
        'komponen_list': komponen_list,
        'form': form
    }
    return render(request, 'dashboard/manajemen_komponen_nilai.html', context)

@user_passes_test(is_admin)
def komponen_nilai_delete(request, pk):
    komponen = get_object_or_404(KomponenNilai, pk=pk)
    if request.method == 'POST':
        komponen.delete()
    return redirect('komponen_nilai_list')

# --- API VIEW BARU UNTUK DAFTAR SISWA ---

def api_get_siswa(request, kelas_id):
    """
    API endpoint untuk mengambil daftar siswa dari sebuah kelas.
    """
    try:
        kelas = ClassGroup.objects.get(pk=kelas_id)
        siswa_list = kelas.students.all().order_by('first_name', 'last_name')
        data = [{'id': siswa.pk, 'name': siswa.get_full_name() or siswa.username} for siswa in siswa_list]
        return JsonResponse(data, safe=False)
    except ClassGroup.DoesNotExist:
        return JsonResponse({'error': 'Kelas tidak ditemukan'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def api_get_rapor_data(request, siswa_id, mapel_id):
    """
    API endpoint untuk mengambil data rapor lengkap, termasuk semua komponen penilaian.
    """
    try:
        # Cari atau buat objek Rapor utama untuk siswa dan mapel ini
        rapor, created = Rapor.objects.get_or_create(
            siswa_id=siswa_id,
            mata_pelajaran_id=mapel_id
        )

        # Ambil semua entri Penilaian yang terkait dengan Rapor ini
        semua_penilaian = Penilaian.objects.filter(rapor=rapor).order_by('tanggal')

        # Ubah data penilaian menjadi format yang mudah dibaca oleh JavaScript
        list_penilaian = [
            {
                'id': p.id,
                'nama_penilaian': p.nama_penilaian,
                'komponen_id': p.komponen.id if p.komponen else None,
                'komponen_nama': p.komponen.nama_komponen if p.komponen else '',
                'nilai': p.nilai
            }
            for p in semua_penilaian
        ]
        
        # Ambil semua jenis komponen yang tersedia untuk dropdown "Tambah Penilaian"
        semua_komponen = KomponenNilai.objects.all().values('id', 'nama_komponen')

        # Siapkan data lengkap untuk dikirim sebagai JSON
        data = {
            'rapor_id': rapor.id,
            'nilai_uts': rapor.nilai_uts,
            'nilai_uas': rapor.nilai_uas,
            'catatan_guru': rapor.catatan_guru,
            'nilai_akhir_rata_rata': rapor.nilai_akhir_rata_rata,
            'penilaian': list_penilaian, # Daftar semua tugas, ulangan, dll.
            'komponen_tersedia': list(semua_komponen) # Untuk form tambah nilai
        }
        try:
            kehadiran_obj = Kehadiran.objects.get(siswa_id=siswa_id)
            kehadiran_json = {'sakit': kehadiran_obj.sakit, 'izin': kehadiran_obj.izin, 'alpha': kehadiran_obj.alpha}
        except Kehadiran.DoesNotExist:
            kehadiran_json = {'sakit': 0, 'izin': 0, 'alpha': 0}
        ekskul_json = [
             {'nama_kegiatan': e.nama_kegiatan, 'predikat': e.predikat or '', 'keterangan': e.keterangan or ''}
             for e in Ekstrakurikuler.objects.filter(siswa_id=siswa_id)
        ]
        prestasi_json = [
            {'jenis_kegiatan': p.jenis_kegiatan, 'keterangan': p.keterangan}
            for p in Prestasi.objects.filter(siswa_id=siswa_id)
        ]
        catatan_wk = (CatatanWaliKelas.objects.filter(siswa_id=siswa_id).first() or None)
        tanggapan_ortu = (TanggapanOrangTua.objects.filter(siswa_id=siswa_id).first() or None)
        kenaikan = (KenaikanKelas.objects.filter(siswa_id=siswa_id).first() or None)
        data.update({
            'kehadiran': kehadiran_json,
            'ekstrakurikuler': ekskul_json,
            'prestasi': prestasi_json,
            'catatan_wali_kelas': catatan_wk.catatan if catatan_wk else '',
            'tanggapan_orang_tua': tanggapan_ortu.tanggapan if tanggapan_ortu else '',
            'kenaikan_kelas': {
                'status': kenaikan.status if kenaikan else '',
                'keterangan': kenaikan.keterangan if kenaikan else ''
            }
        })
        return JsonResponse(data)
        
    except Exception as e:
        logger.error(f"Error di api_get_rapor_data: {e}") 
        return JsonResponse({'error': f'Terjadi kesalahan internal: {str(e)}'}, status=500)
@csrf_exempt
@require_POST
@login_required
@user_passes_test(is_admin)
def api_tambah_penilaian(request):
    try:
        data = json.loads(request.body)
        rapor = Rapor.objects.get(pk=data['rapor_id'])
        komponen = KomponenNilai.objects.get(pk=data['komponen_id'])
        
        Penilaian.objects.create(
            rapor=rapor,
            komponen=komponen,
            nama_penilaian=data['nama_penilaian'],
            nilai=data['nilai']
        )
        return JsonResponse({'status': 'success', 'message': 'Penilaian berhasil ditambahkan.'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

@csrf_exempt
@require_POST
@login_required
@user_passes_test(is_admin)
def api_simpan_semua_perubahan(request, rapor_id):
    """
    Menyimpan:
    - Rapor: nilai_uts, nilai_uas, catatan_guru
    - Penilaian harian (list id+nilai)
    - Kehadiran: sakit, izin, alpha (OneToOne per siswa)
    - Ekskul: list item {id? opsional, nama_kegiatan, predikat, keterangan}
    - Prestasi: list item {id? opsional, jenis_kegiatan, keterangan}
    - Catatan wali kelas: text tunggal (simpan sebagai entri terbaru)
    - Tanggapan orang tua: text tunggal (buat/replace entri terbaru)
    - Kenaikan kelas: {status, keterangan}
    """
    try:
        data = json.loads(request.body)
        rapor = Rapor.objects.select_related('siswa').get(pk=rapor_id)
        siswa = rapor.siswa

        # --- Rapor utama
        rapor.nilai_uts = data.get('nilai_uts') or 0
        rapor.nilai_uas = data.get('nilai_uas') or 0
        rapor.catatan_guru = data.get('catatan_guru', '')
        rapor.save()

        # --- Penilaian harian
        for item in data.get('penilaian', []):
            try:
                p = Penilaian.objects.get(pk=item['id'], rapor=rapor)
                p.nilai = item.get('nilai') or 0
                p.save()
            except Penilaian.DoesNotExist:
                # abaikan bila id tak valid
                continue

        # --- Kehadiran (OneToOne)
        kehadiran_payload = data.get('kehadiran')  # {sakit, izin, alpha}
        if kehadiran_payload:
            Kehadiran.objects.update_or_create(
                siswa=siswa,
                defaults={
                    'sakit': int(kehadiran_payload.get('sakit') or 0),
                    'izin': int(kehadiran_payload.get('izin') or 0),
                    'alpha': int(kehadiran_payload.get('alpha') or 0),
                },
            )

        # --- Ekskul (replace sederhana: hapus lalu buat ulang)
        ekskul_payload = data.get('ekstrakurikuler', [])
        if isinstance(ekskul_payload, list):
            siswa.ekskul_siswa.all().delete()
            for e in ekskul_payload:
                if not e.get('nama_kegiatan'):
                    continue
                Ekstrakurikuler.objects.create(
                    siswa=siswa,
                    nama_kegiatan=e['nama_kegiatan'],
                    predikat=e.get('predikat'),
                    keterangan=e.get('keterangan'),
                )

        # --- Prestasi (replace sederhana)
        prestasi_payload = data.get('prestasi', [])
        if isinstance(prestasi_payload, list):
            siswa.prestasi.all().delete()
            for pr in prestasi_payload:
                if not pr.get('jenis_kegiatan'):
                    continue
                Prestasi.objects.create(
                    siswa=siswa,
                    jenis_kegiatan=pr['jenis_kegiatan'],
                    keterangan=pr.get('keterangan', ''),
                )

        # --- Catatan Wali Kelas (simpan satu terbaru)
        catatan_wk = data.get('catatan_wali_kelas')
        if catatan_wk is not None:
            siswa.catatan_wali_kelas.all().delete()
            if catatan_wk.strip():
                CatatanWaliKelas.objects.create(siswa=siswa, catatan=catatan_wk.strip())

        # --- Tanggapan Orang Tua (optional, simpan satu terbaru)
        tanggapan = data.get('tanggapan_orang_tua')
        if tanggapan is not None:
            siswa.tanggapan_orang_tua.all().delete()
            if tanggapan.strip():
                TanggapanOrangTua.objects.create(siswa=siswa, tanggapan=tanggapan.strip())

        # --- Kenaikan Kelas (status + keterangan)
        kenaikan = data.get('kenaikan_kelas')  # {status, keterangan}
        if kenaikan:
            siswa.kenaikan_kelas.all().delete()
            if kenaikan.get('status'):
                KenaikanKelas.objects.create(
                    siswa=siswa,
                    status=kenaikan['status'],
                    keterangan=kenaikan.get('keterangan', ''),
                )

        return JsonResponse({'status': 'success', 'message': 'Semua perubahan berhasil disimpan.'})
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

@login_required
@user_passes_test(is_admin)
def export_nilai_kelas(request, kelas_id, mapel_id):
    """
    Menghasilkan file EXCEL (.xlsx) berisi nilai semua siswa di kelas tertentu
    untuk satu mata pelajaran.
    """
    try:
        kelas = ClassGroup.objects.get(pk=kelas_id)
        mapel = MataPelajaran.objects.get(pk=mapel_id)
        
        # Pengaturan untuk file Excel
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = f'attachment; filename="nilai_{kelas.name}_{mapel.nama}.xlsx"'

        # Membuat workbook dan worksheet Excel
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = 'Nilai Siswa'

        # Header untuk file Excel
        columns = ['NISN', 'Nama Siswa', 'Komponen Penilaian', 'Nama Penilaian', 'Nilai']
        worksheet.append(columns)

        # Mengisi data ke dalam worksheet
        students = kelas.students.all().order_by('username')
        for student in students:
            try:
                rapor = Rapor.objects.get(siswa=student, mata_pelajaran=mapel)
                
                # Tulis nilai UTS dan UAS
                worksheet.append([student.username, student.get_full_name(), 'UTS', 'Ujian Tengah Semester', rapor.nilai_uts])
                worksheet.append([student.username, student.get_full_name(), 'UAS', 'Ujian Akhir Semester', rapor.nilai_uas])

                # Tulis semua penilaian lainnya (tugas, ulangan, dll.)
                for penilaian in Penilaian.objects.filter(rapor=rapor):
                    worksheet.append([student.username, student.get_full_name(), penilaian.komponen.nama_komponen, penilaian.nama_penilaian, penilaian.nilai])
            except Rapor.DoesNotExist:
                worksheet.append([student.username, student.get_full_name(), 'Data Belum Ada', '', ''])
        
        # Simpan workbook ke response
        workbook.save(response)
        return response

    except Exception as e:
        return HttpResponse(f"Terjadi error: {e}", status=500)


@csrf_exempt
@require_POST
@login_required
@user_passes_test(is_admin)
def import_nilai_kelas(request, kelas_id, mapel_id):
    """
    Memproses file CSV yang diunggah untuk mengimpor nilai.
    """
    if "file_import" not in request.FILES:
        return JsonResponse({'status': 'error', 'message': 'File tidak ditemukan.'}, status=400)
        
    csv_file = request.FILES["file_import"]
    if not csv_file.name.endswith('.csv'):
        return JsonResponse({'status': 'error', 'message': 'Format file harus CSV.'}, status=400)

    try:
        mapel = MataPelajaran.objects.get(pk=mapel_id)
        decoded_file = csv_file.read().decode('utf-8').splitlines()
        reader = csv.DictReader(decoded_file)
        
        errors = []
        success_count = 0
        
        for row in reader:
            nisn = row.get('NISN')
            komponen_nama = row.get('Komponen Penilaian')
            nilai = row.get('Nilai')
            
            if not nisn or not komponen_nama or not nilai:
                errors.append(f"Baris tidak lengkap: {row}")
                continue

            try:
                student = User.objects.get(username=nisn)
                rapor, _ = Rapor.objects.get_or_create(siswa=student, mata_pelajaran=mapel)
                
                if komponen_nama == 'UTS':
                    rapor.nilai_uts = nilai
                    rapor.save()
                    success_count += 1
                elif komponen_nama == 'UAS':
                    rapor.nilai_uas = nilai
                    rapor.save()
                    success_count += 1
                else:
                    nama_penilaian = row.get('Nama Penilaian', komponen_nama) # Fallback
                    komponen, _ = KomponenNilai.objects.get_or_create(nama_komponen=komponen_nama)
                    Penilaian.objects.update_or_create(
                        rapor=rapor, komponen=komponen, nama_penilaian=nama_penilaian,
                        defaults={'nilai': nilai}
                    )
                    success_count += 1
            except User.DoesNotExist:
                errors.append(f"Siswa dengan NISN '{nisn}' tidak ditemukan.")
            except Exception as e:
                errors.append(f"Error pada baris untuk NISN '{nisn}': {e}")

        if errors:
            return JsonResponse({'status': 'warning', 'message': f'{success_count} data berhasil diimpor, tetapi ada beberapa error.', 'errors': errors})
        
        return JsonResponse({'status': 'success', 'message': f'Semua {success_count} data berhasil diimpor.'})

    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Terjadi kesalahan saat memproses file: {e}'}, status=500)

@login_required
def exam_preview_view(request, pk):
    exam = get_object_or_404(Exam, pk=pk)
    questions = exam.questions.all().order_by('id')
    
    # =======================================================
    # PERUBAHAN: Konversi data soal ke format JSON
    # =======================================================
    questions_list = []
    for q in questions:
        questions_list.append({
            'id': q.id,
            'text': q.question_text,
            'image_url': q.question_image.url if q.question_image else None,
            'type': q.get_question_type_display(),
            'options': q.options,
        })
    
    questions_json = json.dumps(questions_list, cls=DjangoJSONEncoder)

    context = {
        'exam': exam,
        'questions_json': questions_json, # Kirim data JSON
    }
    return render(request, 'dashboard/exam_preview.html', context)
@login_required
def grade_report_view(request):
    # Ambil semua percobaan ujian yang sudah selesai,
    # yang ujiannya dibuat oleh guru yang sedang login.
    attempts = ExamAttempt.objects.filter(
        exam__created_by=request.user, 
        is_completed=True
    ).select_related('student', 'exam', 'student__class_group').order_by('student__class_group__name', 'exam__name')

    # Gunakan defaultdict untuk pengelompokan yang lebih mudah
    grades_by_class = defaultdict(list)

    for attempt in attempts:
        # Asumsi: Anda punya model UserProfile yang terhubung ke User
        # dan UserProfile punya ForeignKey ke ClassGroup (model kelas)
        # Sesuaikan path ini jika struktur model Anda berbeda
        try:
            class_name = attempt.student.userprofile.class_group.name
        except AttributeError:
            class_name = "Tanpa Kelas" # Default jika siswa tidak punya kelas
        
        grades_by_class[class_name].append(attempt)

    context = {
        'grades_by_class': dict(grades_by_class), # Konversi kembali ke dict biasa
        'page_title': 'Laporan Nilai Ujian'
    }
    return render(request, 'dashboard/grade_report.html', context)
@login_required
def class_student_list_view(request, class_id):
    # Ambil data kelas berdasarkan ID dari URL
    class_group = get_object_or_404(ClassGroup, pk=class_id)
    
    # Ambil semua pengguna yang memiliki peran 'SISWA' dan berada di kelas tersebut
    students = User.objects.filter(
        class_group=class_group,
        role='SISWA'
    ).order_by('first_name', 'last_name')

    context = {
        'class_group': class_group,
        'students': students,
    }
    return render(request, 'dashboard/class_student_list.html', context)

@login_required
def save_exam_progress(request, attempt_pk):
    if request.method == 'POST':
        try:
            attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
            if attempt.is_completed:
                return JsonResponse({'status': 'error', 'message': 'Ujian sudah selesai'}, status=400)

            data = json.loads(request.body)
            attempt.current_answers = data.get('answers', {})
            attempt.save()
            return JsonResponse({'status': 'success', 'message': 'Progress disimpan'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)
@login_required
def submit_exam_view(request, attempt_pk):
    if request.method == 'POST':
        try:
            attempt = get_object_or_404(ExamAttempt, pk=attempt_pk, student=request.user)
            if attempt.is_completed:
                return JsonResponse({'status': 'error', 'message': 'Ujian sudah pernah disubmit.'}, status=400)

            data = json.loads(request.body)
            final_answers = data.get('answers', {})
            
            # Lakukan proses penilaian di sini (Anda perlu membuat fungsi ini)
            # score, total_correct, results_detail = calculate_score(attempt.exam, final_answers)
            
            attempt.final_answers = final_answers
            attempt.is_completed = True
            attempt.end_time = timezone.now()
            # attempt.score = score 
            attempt.save()
            
            redirect_url = reverse('exam_result_detail', kwargs={'attempt_pk': attempt.pk})
            return JsonResponse({'status': 'success', 'message': 'Ujian berhasil dikumpulkan!', 'redirect_url': redirect_url})

        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)
def print_exam_cards_view(request, class_pk):
    class_group = get_object_or_404(ClassGroup, pk=class_pk)
    
    # PERBAIKAN: Menggunakan relasi langsung dari User ke ClassGroup
    # 'studentprofile__class_group' diganti menjadi 'class_group'
    students = User.objects.filter(class_group=class_group)

    context = {
        'class_group': class_group,
        'students': students,
    }
    return render(request, 'dashboard/print_exam_cards.html', context)
@login_required
def select_class_for_printing_view(request):
    # Asumsi guru hanya bisa melihat kelas yang mereka ajar.
    # Sesuaikan filter ini dengan logika di aplikasi Anda.
    class_groups = ClassGroup.objects.all()
    
    if request.method == 'POST':
        class_pk = request.POST.get('class_group')
        if class_pk:
            return redirect('print_exam_cards', class_pk=class_pk)
    
    context = {
        'class_groups': class_groups,
    }
    return render(request, 'dashboard/select_class_for_printing.html', context)
@login_required
def student_profile_edit_view(request, student_pk):
    student = get_object_or_404(User, pk=student_pk)
    if request.method == 'POST':
        form = StudentProfileForm(request.POST, request.FILES, instance=student)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profil siswa berhasil diperbarui.')
            # Arahkan kembali ke halaman daftar siswa di kelasnya
            return redirect('class_student_list', pk=student.class_group.pk)
    else:
        form = StudentProfileForm(instance=student)
    
    context = {
        'form': form,
        'student': student
    }
    return render(request, 'dashboard/student_profile_edit.html', context)